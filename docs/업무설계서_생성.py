# -*- coding: utf-8 -*-
"""수수료 이벤트 정형화 플랫폼 — 업무 설계서(v0.6). 개조식·맑은 고딕·검정 제목·다이어그램."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = '맑은 고딕'
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x55, 0x55, 0x55)
DIR = '/Users/yujin-an/dev/fees/docs/diagrams'

doc = Document()
st = doc.styles['Normal']
st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def set_ko(run):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:ascii'), FONT)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def h(text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text); set_ko(r); r.font.color.rgb = BLACK
    return p

def para(text, color=None, italic=False, size=10.5, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text); set_ko(r); r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p

def b(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    # 개조식: 굵은 머리어 지원(‘머리 :: 설명’)
    if ' :: ' in text:
        head, rest = text.split(' :: ', 1)
        r = p.add_run(head); set_ko(r); r.bold = True
        r2 = p.add_run(' — ' + rest); set_ko(r2)
    else:
        r = p.add_run(text); set_ko(r)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(t.rows[0].cells):
        c.text = ''; r = c.paragraphs[0].add_run(headers[i]); set_ko(r); r.bold = True; r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(cells):
            c.text = ''; r = c.paragraphs[0].add_run(str(row[i])); set_ko(r); r.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def pic(name, w=16.6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f'{DIR}/{name}.png', width=Cm(w))
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)

# ------------------------------------------------------------ 표지
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('수수료 이벤트 정형화 플랫폼'); set_ko(r); r.bold = True; r.font.size = Pt(24)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('업무 설계서'); set_ko(r); r.font.size = Pt(16); r.font.color.rgb = GREY
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run('버전 0.6  ·  2026-07-05  ·  업무개발팀 / 현업 공유용'); set_ko(r); r.font.size = Pt(10.5); r.font.color.rgb = GREY
doc.add_paragraph()
b('목적 :: 수수료 우대(이벤트·협의)의 등록~체결 부과 전 과정을 업무 흐름·구조 중심으로 정리')
b('대상 :: 현업(마케팅·PB)·업무개발 담당자 (개발 용어 배제)')
b('구성 :: 업무 프로세스 / 데이터 구조 / 적용 방식 / 아키텍처 / 체결 적용 플로우')

h('용어', level=2)
table(['용어', '뜻'], [
    ['기본수수료', '전 계좌 적용 상품군별 표준요율(우대 비교 기준선)'],
    ['이벤트수수료', '한시적 우대(일괄·신청·가입·휴면복귀로 대상 결정)'],
    ['협의수수료', '개별 고객 우대(자격 조건 충족 시)'],
    ['적용범위', '우대 적용 축 — 상품군·거래소·세션·채널(파생=품목)'],
    ['요율표', '수수료 구성 항목표(자사·유관기관·세금)'],
    ['해석', '체결 시점에 계좌별 최저 수수료를 즉석 산출'],
], widths=[3.5, 12])
doc.add_page_break()

# ============================================================ 1
h('1. 업무 프로세스', level=1)

h('1.1 개요', level=2)
b('현업이 우대를 규칙(룰)으로 등록 → 심사·승인 → 활성')
b('고객 체결 순간, 계좌에 가장 유리한(최저) 수수료 자동 선택')
b('핵심 :: 계좌 사전 전개 없음 — 규칙만 저장, 체결 시 해석')
pic('d1_process')

h('1.2 수수료 3계층 · 최저가 자동 적용', level=2)
table(['계층', '내용', '동률 우선'], [
    ['기본(BASE)', '상품군별 표준요율(항상 후보)', '3순위'],
    ['이벤트(EVENT)', '한시 우대(4유형)', '2순위'],
    ['협의(NEGOTIATED)', '개별 고객 우대(조건 충족)', '1순위'],
], widths=[3.5, 9, 3])
b('겹치는 후보 중 고객 부담 최저 자동 적용')
b('동률 시 우선순위 :: 협의 > 이벤트 > 기본')

h('1.3 이벤트 대상 편입 4유형', level=2)
table(['유형', '대상 결정', '예시'], [
    ['일괄적용형', '적용범위 내 전 계좌 자동(신청 불필요)', 'NXT 거래 전 계좌 인하'],
    ['신청형', '고객 신청·승인 계좌만', '온라인 약정 우대'],
    ['가입형', '이벤트 가입 계좌만', '신규 가입 우대'],
    ['휴면복귀형', '휴면 복귀 계좌 자동', '복귀 고객 면제'],
], widths=[2.6, 9, 4])

h('1.4 이벤트 기간 — 신청 가능기간 vs 적용기간', level=2)
b('신청/유입 가능기간 :: 계좌가 이벤트에 들어올 수 있는 캘린더 구간')
b('적용기간(혜택) :: 실제 우대 적용 구간 — 룰마다 택1')
b('캘린더 고정형 :: 룰 종료일까지(모두 같은 날 종료)', level=1)
b('유입시점 상대형 :: 계좌별 “가입일 + N개월”', level=1)
b('예: “가입 시 두 달 무료” → 6/20 가입 = 8/20까지, 신청 마감(6/30) 지나도 유지', level=1)
b('상대형은 고객별 종료일 상이 → 체결 시 해석이 자동 반영')
pic('d2_period')

h('1.5 등록 → 승인 → 활성', level=2)
b('① 현업 기안 :: 적용범위·요율표·기간·대상 지정')
b('② 시뮬레이션 검증 :: 지배관계(전 구간 기존요율 이하)·역마진 경고')
b('③ 심사 승인 :: 검토 후 승인/반려')
b('④ 활성 :: 이후 사전 전개 없이 체결 시 해석 반영')
b('주의 :: 주식은 종목 단위 우대 없음 — 거래소·세션·채널로만 결정(시뮬레이션도 구간 단위). 종목 지정은 파생만')

h('1.6 협의수수료 — 신청 · 연장 대상 산출', level=2)
b('신청·승인 :: 계좌별 개별. 자격 조건(예: 6개월 평균자산 5억↑) 충족 시 부여')
b('연장 :: 계좌별이 아니라 그룹 단위로 대상 일괄 산출')
b('그룹 축 :: 주식형 = 상품군 / 파생 = 품목별', level=1)
b('재평가 분류 :: 신규(새로 충족)·유지(계속 충족)·탈락(더는 미충족)', level=1)
b('기존 대비 비교로 “탈락 대상”까지 함께 노출', level=1)
b('수행 :: 만료 임박 자동 산출 → 담당자 확인·조정 → 일괄 승인(정기 배치 아님)')
pic('d7_nego')

h('1.7 휴면복귀 적용', level=2)
b('원장 휴면복귀 감지(시그널) → 휴면복귀형 이벤트 자동 적용(즉시)')

h('1.8 역할', level=2)
table(['역할', '책임'], [
    ['현업(마케팅·PB)', '이벤트·협의 기안, 적용범위·요율표·기간 설계'],
    ['심사', '지배관계·역마진 검토 후 승인/반려'],
    ['업무개발·시스템', '규칙 해석·캐시·무효화, 원장 연계, 지표 적재'],
    ['원장', '체결 처리, 지표 산정, 수수료 실제 부과'],
], widths=[3.5, 12])
doc.add_page_break()

# ============================================================ 2
h('2. 데이터 구조', level=1)
b('물리 테이블 명세가 아니라 “무엇을 어디에 저장하고 어떻게 연결되나” 중심')
pic('d3_data')

h('2.1 핵심 정보 묶음', level=2)
table(['정보', '내용', '비고'], [
    ['요율표', '자사·유관기관·세금 항목과 요율 방식', '이벤트·협의·기본이 참조'],
    ['규칙(룰)', '유형·상태·적용범위·기간(2축)·조건', '이벤트/협의 공통'],
    ['협의 예외', '계좌×요율표·유효기간', '계좌 단위 저장(유일)'],
    ['가입/신청 이력', '계좌×규칙·가입일', '상대형·협의 자격 근거'],
    ['계좌 지표', '6개월 평균자산·약정액', '자격 판정'],
    ['해석 결과 보관', '계좌×조회키 → 답(캐시)', '읽기 지연 감소'],
    ['이력·감사', '규칙 변경·협의 부여/해지·조건 평가', '감사 추적'],
], widths=[3.2, 8, 4.3])

h('2.2 “전량 전개 테이블” 부재 이유', level=2)
b('계좌×종목 사전 전개 = 수천만 계좌 × 수천 종목 → 저장 폭발, 대규모 불가')
b('대안 :: 규칙만 저장 → 체결 시 해석. 계좌 단위 저장은 협의 예외뿐')
doc.add_page_break()

# ============================================================ 3
h('3. 적용 방식 — 배치 / 신청 / 시그널 즉시', level=1)
pic('d4_apply')
table(['업무', '방식', '트리거/주기', '즉시'], [
    ['이벤트 발효·만료', '배치', '기간 도래(일 점검)', '배치 시점'],
    ['일괄적용형 이벤트', '규칙 활성 즉시', '승인·발효', '즉시'],
    ['신청형·가입형 편입', '고객 신청/가입', '고객 행위', '즉시'],
    ['상대형 혜택 종료', '해석 시 판정', '가입일+N 경과', '조회 시'],
    ['휴면복귀 적용', '시그널 즉시', '원장 감지', '즉시'],
    ['계좌 지표 재산정', '배치', '주기(월/일)', '배치 시점'],
    ['협의 자격 평가', '지표 연동', '지표 재산정 후', '갱신 시'],
    ['협의 연장 대상 산출', '요청형', '만료 임박→담당자 확인', '승인 시'],
    ['해석 결과 무효화', '변경 연동', '승인/연장/지표 변경', '즉시'],
], widths=[3.6, 3, 5, 3])
b('정기 배치 필수 :: “기간 도래(발효·만료)” · “지표 재산정”')
b('협의 연장 :: 고정 시각 배치 아님 — 만료 임박 자동 산출 후 담당자 확인(요청형)')
b('신청·가입·휴면복귀·상대형 종료 :: 배치 없이 즉시/조회 시점 반영')
doc.add_page_break()

# ============================================================ 4
h('4. 아키텍처 — 규칙은 저장, 금액은 체결 때 해석', level=1)
pic('d5_arch')

h('4.1 기본 원리', level=2)
b('금액 미저장 :: 규칙만 보관, 체결 시 해석해 즉석 계산')
b('캐시 :: 해석 답을 (계좌, 조회키) 단위 보관 → 읽기 지연 감소')
b('무효화 :: 규칙 변경 시 영향 범위만 무효화')

h('4.2 플랫폼 / 원장 경계', level=2)
table(['영역', '책임'], [
    ['플랫폼', '규칙 등록·심사·승인, 적용범위·기간, 체결 시 해석, 협의 부여/연장, 캐시·무효화'],
    ['원장', '체결·잔고, 계좌 지표 산정·전달, 해석 결과로 실제 부과'],
], widths=[3, 12.5])

h('4.3 규모 대응', level=2)
b('전개형 :: 조회 빠르나 저장 폭증 → 대규모 불가')
b('해석형+캐시 :: 저장은 규칙·예외뿐, 읽기=캐시, 변경=증분 무효화')
b('주식 조회키에서 종목 제거(거래소·세션·채널) → 최대 카디널리티 제거. 파생만 품목 유지')

h('4.4 조회키', level=2)
table(['상품군', '조회키'], [
    ['국내·해외주식, 금현물', '상품군·거래소·세션·채널 (종목 없음)'],
    ['국내·해외파생', '상품군·거래소·세션·채널·품목'],
], widths=[4.5, 11])

h('4.5 성능 — 캐시 미스는 느리지 않은가?', level=2)
b('현행(등급 방식) :: 계좌 등급 조회 → (등급,상품) 요율 조회 = 2회 색인 읽기. 단, 이벤트·협의·최저가·기간 표현 불가')
b('전량 전개 방식 :: 조회 1회로 빠르나 저장 폭발 + 규칙 변경마다 재전개 부담')
b('해석형 캐시 적중 :: 저장된 답 1건 조회 = 전개형과 동급(1회 읽기)')
b('해석형 캐시 미스 :: 그 (계좌,조회키)의 후보(기본+활성 이벤트/협의 소수)만 메모리 내 최저가 계산')
b('미스 비용 특성 :: 추가 I/O 없음(규칙·요율표는 상주) · 후보 수 소수(수 개) · (계좌,조회키)당 최초 1회만 발생', level=1)
b('완화 :: 활성 계좌×빈출 조회키 캐시 예열, 변경 시 영향분만 무효화 → 정상 운영에서 미스 희소', level=1)
b('결론 :: 미스 1건은 단순 2회 읽기보다 연산은 많으나 I/O가 없고 최초 1회뿐. 반면 등급 방식으로는 이벤트·협의·최저가를 애초에 표현 못 하며, 이를 담으려면 전개(저장 폭발) 또는 어차피 체결 시 최저가 계산이 필요')
doc.add_page_break()

# ============================================================ 5
h('5. 체결 시 수수료 적용 플로우', level=1)
pic('d6_flow', w=13.5)
b('① 체결 발생 :: 계좌·거래소·세션·채널·(파생)품목·체결가·수량 확정')
b('② 조회키 구성 :: 주식은 종목 붕괴 → 거래소·세션·채널')
b('③ 후보 수집 :: 협의·이벤트·기본(편입·적용기간 통과분만)')
b('④ 최저가 결정 :: 고객 부담 최저 선택(동률 협의>이벤트>기본) → 답 캐시 저장')
b('⑤ 금액 산정 :: 요율표 × 체결가·수량 즉석 계산(자사·유관기관·세금)')
b('⑥ 원장 반영 :: 수수료 구성 전달 → 실제 부과·잔고 반영')

h('5.1 캐시 적중/미스', level=2)
b('첫 조회(미스) :: 후보 계산 후 답 저장')
b('이후 조회(적중) :: 저장된 답 즉시 사용(재계산 없음)')
b('변경 시 :: 영향 범위만 무효화 → 다음 조회 때 재해석')

h('5.2 예시 — 상대형 이벤트', level=2)
b('신규 가입 2개월 무료 · 6/20 가입 고객, 7월 초 국내주식 온라인 체결 → 자사 수수료 0원 무료 요율(유관·세금 정상)')
b('4월 가입 고객(2개월 경과) → 무료 아님, 표준요율 복귀')
b('동일 이벤트, 고객별 상이 :: 적용기간이 “가입일 기준 상대형”이기 때문')

out = '/Users/yujin-an/dev/fees/docs/수수료플랫폼_업무설계서_v0.6.docx'
doc.save(out)
print('saved', out)
