# 기술설계서 v1.8 생성기 — v1.7.docx를 읽어 이번 확정분을 additive 반영해 v1.8.docx로 저장.
#   ① 표지 버전 v1.7 → v1.8
#   ② §10.4 후보 색인 4축 → 6축 정정(자산군·조회구분·거래소·품목 +세션·채널), tie_order·specificity 명시
#   ③ §10.5 신설 — 1차전의 물리화: 색인 키 4축→6축 완전 키화 + 계산을 인메모리에서 물리 테이블로
#   ④ 프로세스 도식 삽입(docs/diagrams/v18_1차전_프로세스.png, 한국식 컨설팅 스타일)
#   ⑤ 판 이력 갱신
# 체인: v1.3(진본) → v14 → v15 → v16 → v17 → v18. 실행: python3 docs/scripts/기술설계서_v18_생성.py (레포 루트 기준)
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = 'docs/수수료정책플랫폼_기술설계서_v1.7.docx'
DST = 'docs/수수료정책플랫폼_기술설계서_v1.8.docx'
IMG = 'docs/diagrams/v18_1차전_프로세스.png'


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def set_para_text(p, text):
    """단락의 모든 런을 비우고 첫 런에 텍스트를 몰아넣어 스타일을 보존한다."""
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def main():
    d = Document(SRC)

    # ① 표지 버전 (런이 '버전'+' v1.7'로 갈라져 있어 단락 단위로 찾는다)
    for p in d.paragraphs:
        if '버전 v1.7' in p.text:
            for r in p.runs:
                if 'v1.7' in r.text:
                    r.text = r.text.replace('v1.7', 'v1.8')

    # 목록(글머리) 스타일 견본 — 기존 a0 단락의 스타일을 빌린다
    bullet_style = next(p.style for p in d.paragraphs
                        if p.style.style_id == 'a0' and p.text.strip())

    # ② §10.4 후보 색인 4축 → 6축 정정
    idx_p = next(p for p in d.paragraphs if p.text.startswith('조회키 조합별 후보 색인'))
    set_para_text(idx_p,
        '조회키 조합별 후보 색인: 조합(자산군·조회구분·거래소·품목·세션·채널 6축)마다 순위값 오름차순 후보 '
        '목록을 소형 테이블로 두고, 정책 승인·종료 시 그 정책이 걸리는 조합만 갱신함. 색인 각 행에 '
        'tie_order·specificity를 함께 저장해 조합 횡단 점조회 정렬에 씀. 색인 전체가 메가바이트 규모라 '
        '유지 비용이 없음(축 확장·물리화 상세는 §10.5)')

    # ③·④ §10.5 신설 — "11. 등록" H1 앞
    h11 = next(p for p in d.paragraphs
               if p.style.name == 'Heading 1' and p.text.startswith('11. 등록'))

    h11.insert_paragraph_before(
        '10.5 1차전의 물리화: 색인 키 4축→6축, 계산은 인메모리에서 물리 테이블로', style='Heading 2')
    h11.insert_paragraph_before(
        '배정판 산출은 1차전(계좌 무관 사전 랭킹)과 2차전(계좌 배정)으로 나뉨. 이번에 1차전에서 두 가지가 '
        '확정됨. 하나는 후보 색인의 키를 4축에서 6축으로 넓힌 것이고, 다른 하나는 후보·순위를 배치 실행 중 '
        '인메모리로 계산하던 것을 정책 승인 시점에 물리 테이블로 옮긴 것임. 아래 도식이 등록부터 원장 조회까지의 '
        '전체 흐름과 1차전에서 일어난 두 변화를 함께 보여줌.')

    # 프로세스 도식 삽입 (중앙 정렬)
    pic_p = h11.insert_paragraph_before('')
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(IMG, width=Cm(16.6))
    cap_p = h11.insert_paragraph_before('[그림] 새 수수료 정책의 여정과 1차전의 두 변화(4축→6축·물리화)')
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 가. 색인 키 4축 → 6축 완전 키화
    h11.insert_paragraph_before('가. 색인 키 4축 → 6축 완전 키화')
    for t in [
        '기존 후보 색인은 자산군·조회구분·거래소·품목 4축으로 조합을 만들었음. 세션·채널은 대개 "전체" 한 '
        '칸이라 키에서 빠져 있었음',
        '세션·채널만 한정한 정책(야간 세션 이벤트, MTS 채널 이벤트)은 그 셀의 승자를 바꿈. 이런 정책이 '
        '색인 밖에 있으면 2차전이 계좌 루프를 돌며 범위를 다시 훑어야 했음',
        '그래서 세션·채널을 색인 키로 승격해 6축(자산군·조회구분·거래소·품목·세션·채널)으로 완전 키화함. '
        '수수료 승자를 가르는 계좌 무관 축이 모두 키에 들어옴',
        '효과: 2차전은 계좌×셀 키 조회와 자격 게이트만 남고, 계좌 루프에서 범위 매칭·셀 축 재수집이 사라짐',
    ]:
        h11.insert_paragraph_before(t, style=bullet_style)

    # 나. 계산은 인메모리에서 물리 테이블로
    h11.insert_paragraph_before('나. 계산은 인메모리에서 물리 테이블로')
    for t in [
        '기존에는 후보 목록과 순위를 배치가 돌 때마다 인메모리로 다시 계산했음(계좌×셀마다 후보를 재수집). '
        '비용이 배치 실행에 실렸음',
        '이제는 정책 승인·종료 트랜잭션에서 두 가지를 물리 테이블로 확정함: (1) 규칙 행에 순위값 스탬프, '
        '(2) 6축 조합별 후보 색인 테이블(각 행에 tie_order·specificity 포함)',
        '이후 모든 경로(일배치 전체 재산출·수시 증분·화면 조회·미스 경로)는 저장된 순위값을 읽기만 함. '
        '우선순위 판단은 승인 시점 한 곳에만 존재함',
        '색인은 활성 정책 수(수백 건)에 비례해 메가바이트 규모라 유지 비용이 없고, 갱신은 그 정책이 걸리는 '
        '조합만 대상으로 함',
    ]:
        h11.insert_paragraph_before(t, style=bullet_style)

    h11.insert_paragraph_before(
        '이 물리화로 지배 비용은 결과 정렬·적재로 수렴함(§10.3). 1차전은 승인 때 한 번, 2차전은 배치마다 '
        '돌지만 둘 다 같은 저장 순위값을 읽으므로 두 경로의 결과가 항상 일치함.')

    # ⑤ 판 이력 갱신
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.startswith('v1.0 기술설계서'):
                    c.text = ('v1.0 기술설계서 → v1.1·v1.2 우선순위 저장 설계 → v1.3 전체 통합·한글화 → '
                              'v1.4~v1.6 배치 전략·배정판 저장 범위 → v1.7 편입내역 용어 정비·셀 전개 곱집합·'
                              '수십억 후보·우선순위 처리 → v1.8(본 문서, 1차전 색인 4축→6축 완전 키화·'
                              '인메모리에서 물리 테이블화·프로세스 도식)')

    d.save(DST)

    # 검증
    d2 = Document(DST)
    body = '\n'.join(p.text for p in all_paragraphs(d2))
    print(f'저장: {DST} / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)}')
    print('표지:', next(p.text for p in d2.paragraphs if '버전 v' in p.text))
    checks = [
        '자산군·조회구분·거래소·품목·세션·채널 6축',
        '10.5 1차전의 물리화',
        '가. 색인 키 4축 → 6축 완전 키화',
        '나. 계산은 인메모리에서 물리 테이블로',
        'tie_order·specificity',
        'v1.8(본 문서',
    ]
    for key in checks:
        print(f'  {"OK" if key in body else "누락!!"} — {key}')
    imgs = len(d2.inline_shapes)
    print(f'인라인 이미지: {imgs}개')


if __name__ == '__main__':
    main()
