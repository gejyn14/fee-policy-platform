# 기술설계서 v1.9 생성기 — v1.8.docx를 읽어 두 가지를 반영해 v1.9.docx로 저장.
#   (1) 8축: 유동성(메이커/테이커)·매수매도를 조회키 축으로 승격 → 색인 키 6축 → 8축
#       (§조회키 구성·축 배치 원칙·후보 색인·1차전 물리화·프로세스 도식 전부 8축으로 정정)
#   (2) 문서 재구성(프로세스 전진 배치): 새 장 '정책의 여정'을 앞에 두고, 1차전 경쟁(우선순위)을
#       바로 뒤로, 데이터 모델은 뒤로 이동. 전체 장 번호 재정렬 + 목차·상호참조(§N) 일괄 정정.
#   장 번호 재배치: 0·1 유지, [신규 2=정책의 여정], 구5→3, 구3→4, 구2→5, 구4→6, 구6→7,
#                   구7→8, 구8→9, 구9→10, 구10→11, 구11→12, 부록 유지.
# 체인: v1.3(진본) → v14 → … → v18 → v19. 실행: python3 docs/scripts/기술설계서_v19_생성.py
import re
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = 'docs/수수료정책플랫폼_기술설계서_v1.8.docx'
DST = 'docs/수수료정책플랫폼_기술설계서_v1.9.docx'
IMG = 'docs/diagrams/v19_정책여정_프로세스.png'

# 구 장 정수 → 신 장 정수 (신규 2번은 별도 삽입)
CHMAP = {0: 0, 1: 1, 2: 5, 3: 4, 4: 6, 5: 3, 6: 7, 7: 8, 8: 9, 9: 10, 10: 11, 11: 12}
# 상호참조 토큰 정정(구 번호 → 신 번호). §1.4는 불변이라 생략.
REFMAP = {'§5.1': '§3.1', '§10.1': '§11.1', '§10.2': '§11.2', '§10.3': '§11.3', '§10.5': '§11.5'}

NEW_TOC = [
    '0. 수수료 정책 플랫폼 개요',
    '1. 아키텍처: 정책은 저장, 배정판은 사전 산출',
    '2. 정책의 여정: 등록에서 원장 조회까지',
    '3. 수수료 우선순위(1차전 경쟁 규칙): 세 정책이 한 순위에서 경쟁',
    '4. 조회키: 체결을 수수료 정책에 잇는 기준',
    '5. 데이터 모델: 무엇을 어디에 담는가',
    '6. 수수료 계산: 요율표 평가',
    '7. 개별 체결의 수수료 결정: 원장의 배정판 조회',
    '8. 대상 편입: 고객이 정책 대상이 되는 방식',
    '9. 이벤트 기간: 신청 가능기간과 적용기간',
    '10. 협의수수료: 신청·자격·승인·연장',
    '11. 배정판 산출·무효화: 일배치 + 수시 증분',
    '12. 등록 → 검증 → 승인 → 활성',
    '부록. 배정판 점조회 성능 실측',
]


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def set_text(p, text):
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def find(doc, pred):
    return next(p for p in all_paragraphs(doc) if pred(p))


def replace_token_everywhere(doc, old, new):
    """모든 단락에서 토큰 치환. 런 경계에 걸리면 그 단락만 평탄화."""
    for p in all_paragraphs(doc):
        if old not in p.text:
            continue
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
        if old in p.text:
            set_text(p, p.text.replace(old, new))


def main():
    d = Document(SRC)

    # ── 스타일 견본
    h1_style = next(p.style for p in d.paragraphs if p.style.name == 'Heading 1')
    normal_style = next(p.style for p in d.paragraphs
                        if p.style.name == 'Normal' and p.text.strip())
    bullet_style = next(p.style for p in d.paragraphs
                        if p.style.style_id == 'a0' and p.text.strip())
    num_style = next(p.style for p in d.paragraphs if p.style.name == 'List Number')
    toc_style = None  # 목차 항목 스타일(뒤에서 채움)

    # ── 0) 표지 버전 v1.8 → v1.9
    for p in d.paragraphs:
        if '버전 v1.8' in p.text:
            for r in p.runs:
                if 'v1.8' in r.text:
                    r.text = r.text.replace('v1.8', 'v1.9')

    # ── 1) 8축 내용 정정 (장 번호 정정 前, 원문 텍스트 기준)
    # (a) §조회키 구성: 체결 정보 목록에 유동성·매수매도 추가
    comp = find(d, lambda p: p.text.startswith('체결이 들어오면 원장에 함께 들어오는 정보'))
    set_text(comp, comp.text.replace(
        '(자산군·거래소·조회구분·세션·품목·채널)',
        '(자산군·거래소·조회구분·세션·품목·채널·유동성·매수매도 8축)'))

    # (b) §축 배치 원칙: 매수/매도 승격 + 유동성 축 신설
    sell = find(d, lambda p: p.text.startswith('매수/매도: 요율표 내부 항목'))
    set_text(sell, '매수/매도: 판의 키(“전체” 허용). 한쪽만 우대하는 정책'
                   '(매수 무료·매도 면제)이 그 방향의 승자를 바꿈')
    sell.insert_paragraph_before(
        '유동성(메이커/테이커): 판의 키(“전체” 허용). 메이커 우대처럼 유동성 구분을 한정한 '
        '이벤트가 승자를 바꿈(파생·해외 체결에 메이커/테이커 구분이 실려 옴)', style=bullet_style)

    # (c) §후보 색인: 6축 → 8축
    idxb = find(d, lambda p: p.text.startswith('조회키 조합별 후보 색인'))
    set_text(idxb, idxb.text.replace(
        '자산군·조회구분·거래소·품목·세션·채널 6축',
        '자산군·조회구분·거래소·품목·세션·채널·유동성·매수매도 8축'))

    # (d) §1차전 물리화 소절: 제목·본문 6축 → 8축, 도식은 §2로 이관하므로 제거
    h_phys = find(d, lambda p: p.text.startswith('10.5 1차전의 물리화'))
    set_text(h_phys, '10.5 1차전의 물리화: 색인 키 8축 완전 키화, 계산은 물리 테이블로')

    phys_intro = find(d, lambda p: p.text.startswith('배정판 산출은 1차전(계좌 무관 사전 랭킹)과 2차전'))
    set_text(phys_intro,
        '배정판 산출은 1차전(계좌 무관 사전 랭킹)과 2차전(계좌 배정)으로 나뉨(전체 흐름은 §2 도식 참조). '
        '1차전에서 두 가지가 자리 잡음. 하나는 후보 색인의 키를 넓혀 온 것(4축 → 6축 → 8축)이고, 다른 '
        '하나는 후보·순위를 배치 실행 중 인메모리로 계산하던 것을 정책 승인 시점에 물리 테이블로 옮긴 것임.')

    # 도식 + 캡션 단락 제거(§2로 이관)
    for p in list(all_paragraphs(d)):
        if 'graphicData' in p._p.xml or p.text.startswith('[그림] 새 수수료 정책의 여정'):
            p._p.getparent().remove(p._p)

    subA = find(d, lambda p: p.text.startswith('가. 색인 키'))
    set_text(subA, '가. 색인 키 확장: 4축 → 6축 → 8축 완전 키화')
    grow = find(d, lambda p: p.text.startswith('그래서 세션·채널을 색인 키로 승격해 6축'))
    set_text(grow,
        'v1.8에서 세션·채널을 승격해 6축으로, v1.9에서 유동성(메이커/테이커)·매수매도를 승격해 8축으로 '
        '완전 키화함. 수수료 승자를 가르는 계좌 무관 축이 모두 키에 들어옴')
    grow.insert_paragraph_before(
        '유동성·매수매도도 승자를 바꿈: 메이커 우대는 메이커 셀만, 매도 면제·매수 무료는 그 방향 셀만 이김. '
        '색인 밖에 두면 2차전이 그 축을 다시 훑어야 하므로 키로 승격함', style=bullet_style)

    tbl2 = find(d, lambda p: p.text.startswith('이제는 정책 승인·종료 트랜잭션에서 두 가지를 물리 테이블로'))
    set_text(tbl2, tbl2.text.replace('6축 조합별 후보 색인 테이블', '8축 조합별 후보 색인 테이블'))

    # (e) ERD 체결 필드 목록에 유동성·매수매도 추가
    for p in all_paragraphs(d):
        if p.text.strip() == '체결(자산군·거래소·조회구분·세션·채널·품목)':
            set_text(p, '체결(자산군·거래소·조회구분·세션·채널·품목·유동성·매수매도)')

    # ── 2) 장 블록 포착 (재정렬用, 구 번호 기준) + 앵커/목차/이력 포착
    body = d.element.body
    sectPr = body.find(qn('w:sectPr'))
    para_of = {p._p: p for p in d.paragraphs}
    children = [c for c in body if c is not sectPr]

    # H1 경계 인덱스
    h1_idx = [i for i, c in enumerate(children)
              if c.tag == qn('w:p') and para_of.get(c) is not None
              and para_of[c].style.name == 'Heading 1']
    h1_idx_sorted = sorted(h1_idx)

    blocks = {}       # 구 장 정수 → [element,...]
    appendix = None
    for k, hi in enumerate(h1_idx_sorted):
        nxt = h1_idx_sorted[k + 1] if k + 1 < len(h1_idx_sorted) else len(children)
        seg = children[hi:nxt]
        txt = para_of[children[hi]].text
        m = re.match(r'(\d+)\.', txt)
        if m:
            blocks[int(m.group(1))] = seg
        elif txt.startswith('부록'):
            appendix = seg

    anchor = blocks[1][-1]   # 1장(아키텍처) 마지막 요소

    # ── 3) 장 번호 정정: 헤딩 앞자리 + 상호참조 §토큰
    for p in all_paragraphs(d):
        if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
            m = re.match(r'(\d+)', p.text)
            if m and int(m.group(1)) in CHMAP:
                new = str(CHMAP[int(m.group(1))])
                set_text(p, new + p.text[m.end(1):])
    for old, new in sorted(REFMAP.items(), key=lambda kv: -len(kv[0])):
        replace_token_everywhere(d, old, new)

    # 3-특례: 구5(우선순위) 제목에 '1차전 경쟁 규칙' 부기
    tp = find(d, lambda p: p.text.startswith('3. 수수료 우선순위: 세 정책이 한 순위에서 경쟁'))
    set_text(tp, '3. 수수료 우선순위(1차전 경쟁 규칙): 세 정책이 한 순위에서 경쟁')

    # ── 4) 신규 2장 '정책의 여정' 생성(문서 끝에 붙였다가 이동)
    new_elems = []

    def add(text, style):
        p = d.add_paragraph(text, style=style)
        new_elems.append(p._p)
        return p

    add('2. 정책의 여정: 등록에서 원장 조회까지', h1_style)
    add('수수료 정책 하나가 등록되어 실제 체결에 적용되기까지의 전체 여정임. 정책은 등록과 승인을 거친 뒤 '
        '두 번의 경쟁을 통과함. 1차전은 승인 시점에 계좌와 무관하게 스코프 키별 승자 순위를 정하는 사전 '
        '랭킹이고, 2차전은 그 순위를 계좌에 붙여 배정판을 채우는 단계임. 원장은 완성된 배정판에서 한 줄만 '
        '읽음. 이 장은 전체 흐름을 먼저 보이고, 세부 규칙은 이후 장에서 다룸.', normal_style)
    picp = d.add_paragraph('', style=normal_style)
    picp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picp.add_run().add_picture(IMG, width=Cm(16.6))
    new_elems.append(picp._p)
    capp = add('[그림] 새 수수료 정책의 여정: 등록 → 승인 → 1차전(8축 사전 랭킹·물리화) → 2차전 → 원장 조회',
               normal_style)
    capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add('단계별로 보면 다음과 같음.', normal_style)
    for t in [
        '등록: 정책 초안에 요율표를 연결하고 적용범위·기간·대상 편입방식을 지정함(상세 §12)',
        '검증·승인: 지배관계·역마진을 심사하고 자격기준과 유효기간을 확정함(상세 §12)',
        '1차전(승인 때 1회): 계좌와 무관한 스코프 키에서 정책들이 한 순위로 경쟁해 승자 순위를 정함. '
        '순위값과 8축 조합별 후보 색인을 물리 테이블로 저장함(경쟁 규칙 §3, 조회키 8축 §4, 물리화 §11.5)',
        '2차전(배치마다): 계좌×셀에서 저장된 순위를 조회하고 자격 게이트를 적용해 셀별 승자를 확정한 뒤 '
        '배정판(우대분만)을 적재함(§11)',
        '원장 조회(체결 시): 완성된 배정판을 점조회하고, 맞는 행이 없으면 기본수수료를 직접 적용함(§7)',
    ]:
        add(t, num_style)
    add('핵심은 무거운 판단인 우선순위 경쟁을 승인 시점의 1차전 한 곳으로 모으고, 배치마다 도는 2차전은 '
        '저장된 순위를 읽기만 하게 한 것임. 두 경쟁이 같은 저장 순위값을 쓰므로 결과가 항상 일치함.',
        normal_style)

    # ── 5) 블록 재정렬: 1장 뒤에 [신규2, 구5(3), 구3(4), 구2(5), 구4(6), 구6(7), …, 구11(12), 부록]
    order = new_elems + [e for oldch in [5, 3, 2, 4, 6, 7, 8, 9, 10, 11]
                         for e in blocks[oldch]] + appendix
    cur = anchor
    for e in order:
        cur.addnext(e)
        cur = e

    # ── 6) 목차 재작성 (목차 H1 ~ 다음 H1[='0. …'] 사이 항목을 새 목록으로)
    toc_h1 = find(d, lambda p: p.text.strip() == '목차')
    # 목차 항목 단락 수집
    toc_items = []
    seen = False
    for p in d.paragraphs:
        if p._p is toc_h1._p:
            seen = True
            continue
        if seen:
            if p.style.name == 'Heading 1':
                break
            if p.text.strip():
                toc_items.append(p)
    toc_style = toc_items[0].style
    zero_h1 = find(d, lambda p: p.style.name == 'Heading 1' and p.text.startswith('0. 수수료'))
    for i, title in enumerate(NEW_TOC):
        if i < len(toc_items):
            set_text(toc_items[i], title)
        else:
            np = zero_h1.insert_paragraph_before(title, style=toc_style)  # noqa: F841

    # ── 7) 판 이력 갱신
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.startswith('v1.0 기술설계서'):
                    c.text = (
                        'v1.0 기술설계서 → v1.1·v1.2 우선순위 저장 설계 → v1.3 전체 통합·한글화 → '
                        'v1.4~v1.6 배치 전략·배정판 저장 범위 → v1.7 편입내역 용어 정비·셀 전개 곱집합·'
                        '수십억 후보·우선순위 처리 → v1.8 1차전 색인 6축 완전 키화·물리 테이블화·프로세스 도식 → '
                        'v1.9(본 문서, 유동성·매수매도 추가로 8축 완전 키화·문서 재구성[프로세스·1차전 경쟁 '
                        '전진 배치])')

    d.save(DST)

    # ── 검증
    d2 = Document(DST)
    heads = [p.text for p in d2.paragraphs if p.style.name == 'Heading 1']
    print(f'저장: {DST} / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)} / 이미지 {len(d2.inline_shapes)}')
    print('표지:', next(p.text for p in d2.paragraphs if '버전 v' in p.text))
    print('H1 순서:')
    for h in heads:
        print('   ', h[:46])
    body2 = '\n'.join(p.text for p in all_paragraphs(d2))
    for key in ['·유동성·매수매도 8축', '유동성(메이커/테이커): 판의 키',
                '2. 정책의 여정', '색인 키 8축 완전 키화', '8축 조합별 후보 색인 테이블',
                '§11.5', '§3.1', 'v1.9(본 문서']:
        print(f'  {"OK" if key in body2 else "누락!!"} — {key}')
    print('잔존 6축:', body2.count('6축'), '/ 8축:', body2.count('8축'))
    print('잔존 §5.1(0이어야):', body2.count('§5.1'), '/ §10.(0이어야):', len(re.findall(r'§10\.', body2)))


if __name__ == '__main__':
    main()
