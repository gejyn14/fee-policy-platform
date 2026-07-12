# -*- coding: utf-8 -*-
# 개선 제안서 생성기 — 수기 편집본 v2.0.docx(진본)를 읽어 제안서로 재구성해 저장.
#   ① 재프레이밍: 표지를 제안서로 개제, 핵심 요약 → 제안 개요(추진 배경·문제점·개선 방안·기대 효과 표)
#   ② 정합 교정: 수기 편집이 남긴 어긋남 정리(§1.4 잔존 참조, 3층→두 층, 계좌번호 통일, 꼬리 역슬래시 등)
#   ③ 한국식 타이포그래피: 맑은 고딕 통일, 대제목 16 / 중제목 13 / 소제목 11 / 본문 10.5 / 표 9.5pt,
#      남색 제목·대제목 밑줄, 본문 줄간격 1.42, 표 머리행 가운데 정렬·짧은 값 가운데 정렬
# 주의: 원본 v2.0.docx는 사용자가 직접 손본 진본이므로 절대 덮어쓰지 않는다. 산출물만 새로 만든다.
# 실행: python3 docs/scripts/개선제안서_생성.py (레포 루트 기준)
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = 'docs/수수료정책플랫폼_기술설계서_v2.0.docx'
DST = 'docs/수수료정책플랫폼_개선제안서_v1.0.docx'

KFONT = '맑은 고딕'
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
NAVY2 = RGBColor(0x2E, 0x4E, 0x75)
BODY = RGBColor(0x26, 0x26, 0x2B)
GREY = RGBColor(0x60, 0x60, 0x66)
LABEL_FILL = 'EEF2F8'  # 제안 개요 표 왼쪽 칸


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def set_text(p, text):
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def kfont(run, size=None, bold=None, color=None):
    """맑은 고딕(동아시아 포함) + 크기·굵기·색."""
    run.font.name = KFONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), KFONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def h1_rule(p):
    """대제목 밑에 남색 밑줄."""
    ppr = p._p.get_or_add_pPr()
    old = ppr.find(qn('w:pBdr'))
    if old is not None:
        ppr.remove(old)
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F3A5F')
    pbdr.append(bottom)
    ppr.append(pbdr)


def shade(cell, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def main():
    d = Document(SRC)

    # ── 1) 정합 교정 + 가벼운 윤문 (수기 편집 존중, 어긋난 곳만)
    fixes = [
        ('가장 유리한 수수료가 자동으로 적용되도록 하는 시스템임',
         '가장 유리한 수수료를 자동으로 적용하는 시스템임'),
        ('관리해  단일', '관리해 단일'),
        ('§1.4', '§1.2'),                      # 1장 절 삭제로 남은 옛 참조
        ('아래 3층 체제로 전환함', '아래 두 층으로 전환함'),  # ③층 삭제 반영
        ('담당함. \\', '담당함.'),
        ('8041-2237-01', '6041-2237-10'),      # §3.3 수기 변경과 통일
        ('아래 본문은 이 요약을 개요 → 아키텍처 → 데이터 → 결정 흐름 → 정책 유형 → 운영 순으로 자세히 풀어 씀.',
         '아래 본문은 이 제안을 개요 → 아키텍처 → 정책의 여정 → 우선순위 → 조회키 → 대규모 산출 → '
         '데이터 모델 → 운영 순으로 자세히 풀어 씀.'),
    ]
    for old, new in fixes:
        for p in all_paragraphs(d):
            if old in p.text:
                set_text(p, p.text.replace(old, new))

    # ── 2) 표지 개제
    title_p = next(p for p in d.paragraphs if p.text.strip() == '수수료 정책 플랫폼')
    set_text(title_p, '수수료 체계 개선 제안서')
    sub_p = next(p for p in d.paragraphs if p.text.strip() == '기술 설계서')
    set_text(sub_p, '수수료 정책 플랫폼 구축안')
    ver_p = next((p for p in d.paragraphs if p.text.strip() == '버전 v2.0'), None)
    if ver_p is not None:
        ver_p._p.getparent().remove(ver_p._p)
    date_p = next(p for p in d.paragraphs if '임원 보고용' in p.text)
    set_text(date_p, '2026년 7월')

    # ── 3) 핵심 요약 → 제안 개요 (불릿 3개를 4단 개요 표로)
    h_sum = next(p for p in d.paragraphs if p.text.startswith('핵심 요약'))
    set_text(h_sum, '제안 개요')
    kills = [p for p in d.paragraphs
             if p.text.startswith(('왜 필요한가:', '무엇이 핵심인가:', '규모 대응:'))]
    anchor_p = kills[0]._p.getprevious()   # 요약 도입 문장 뒤에 표를 둠
    for p in kills:
        p._p.getparent().remove(p._p)

    rows = [
        ('추진 배경',
         '현행 등급제(계좌 등급 → 등급·상품 요율)는 이벤트·협의·최저가 보장·기간 한정을 표현하지 못함. '
         '마케팅·영업의 다양한 수수료 정책을 담으려면 새로운 틀이 필요함.'),
        ('현황과 문제점',
         '수수료 우대를 정책으로 담을 그릇이 없고, 대안으로 흔히 쓰는 "계좌마다 모든 종목을 미리 펼쳐 '
         '저장"하는 방식은 수백억 건이 되어 성립하지 않음.'),
        ('개선 방안',
         '기본·이벤트·협의를 하나의 "정책"으로 통합 관리함. 정책만 저장하고 매일 배치로 계좌별 수수료 '
         '배정판을 미리 산출하며, 원장은 체결 시 배정판 한 줄만 조회함. 주식은 종목 축을 제거하고 파생만 '
         '품목별로 관리해 단일 테이블로 감당함.'),
        ('기대 효과',
         '고객에게 가장 유리한 수수료가 자동 적용됨. 배정판은 우대분만 저장해 억 단위 대신 수백만 규모에 '
         '머물고(§1.2), 체결 시점 계산이 없어 점조회 1ms 이내를 실측으로 확인함(부록). 배정 이력으로 '
         '민원 대응 추적이 가능함(§6.4).'),
    ]
    ref = d.tables[2]._tbl  # 테두리 서식 견본
    tbl = d.add_table(rows=len(rows), cols=2)
    ref_pr = ref.find(qn('w:tblPr'))
    if ref_pr is not None:
        old_pr = tbl._tbl.find(qn('w:tblPr'))
        if old_pr is not None:
            tbl._tbl.remove(old_pr)
        tbl._tbl.insert(0, copy.deepcopy(ref_pr))
    for i, (label, body) in enumerate(rows):
        lc, bc = tbl.rows[i].cells
        lc.text = label
        bc.text = body
        lc.width, bc.width = Cm(3.2), Cm(13.4)
        shade(lc, LABEL_FILL)
        for r in lc.paragraphs[0].runs:
            kfont(r, 10, True, NAVY)
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in bc.paragraphs[0].runs:
            kfont(r, 9.5, False, BODY)
    anchor_p.addnext(tbl._tbl)

    # ── 4) 타이포그래피 — 스타일 기본값
    for name, sz, bold, col in [('Heading 1', 16, True, NAVY), ('Heading 2', 13, True, NAVY2),
                                ('Normal', 10.5, None, None)]:
        st = d.styles[name]
        st.font.name = KFONT
        st.font.size = Pt(sz)
        if bold:
            st.font.bold = True
        if col:
            st.font.color.rgb = col
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.append(rf)
        rf.set(qn('w:eastAsia'), KFONT)

    # ── 5) 타이포그래피 — 본문 단락 (표 밖)
    toc_zone = False
    for p in d.paragraphs:
        txt = p.text.strip()
        if 'graphicData' in p._p.xml:      # 도식은 손대지 않음
            continue
        if p.style.name == 'Heading 1':
            for r in p.runs:
                kfont(r, 16, True, NAVY)
            h1_rule(p)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            toc_zone = (txt == '목차')
            continue
        if p.style.name == 'Heading 2':
            for r in p.runs:
                kfont(r, 13, True, NAVY2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            continue
        if txt == '수수료 체계 개선 제안서':
            for r in p.runs:
                kfont(r, 27, True, NAVY)
            continue
        if txt == '수수료 정책 플랫폼 구축안':
            for r in p.runs:
                kfont(r, 15, False, GREY)
            continue
        if txt == '2026년 7월':
            for r in p.runs:
                kfont(r, 11, False, GREY)
            continue
        if txt.startswith('[그림]'):
            for r in p.runs:
                kfont(r, 9, False, GREY)
            continue
        if txt.startswith(('가. ', '나. ', '다. ')):   # 소제목
            for r in p.runs:
                kfont(r, 11, True, NAVY2)
            p.paragraph_format.space_before = Pt(8)
            continue
        # 목차 항목·일반 본문
        for r in p.runs:
            kfont(r, 10.5, None, NAVY if toc_zone else BODY)
        if txt:
            p.paragraph_format.line_spacing = 1.42

    # ── 6) 타이포그래피 — 표 (머리행 서식 보존: 크기·글꼴만)
    for t in d.tables:
        for ri, row in enumerate(t.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.12
                    p.paragraph_format.space_after = Pt(2)
                    short = len(p.text.strip()) <= 12 and '·' not in p.text
                    if ri == 0 or short:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        if ri == 0:
                            kfont(r, 9.5)           # 흰 글자·볼드 유지
                        else:
                            kfont(r, 9.5, None, BODY)

    d.save(DST)

    # ── 검증
    d2 = Document(DST)
    body2 = '\n'.join(p.text for p in all_paragraphs(d2))
    print(f'저장: {DST} / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)} / 이미지 {len(d2.inline_shapes)}')
    checks = ['수수료 체계 개선 제안서', '제안 개요', '추진 배경', '기대 효과',
              '6041-2237-10', '두 층으로 전환함']
    for k in checks:
        print(f'  {"OK" if k in body2 else "누락!!"} — {k}')
    for bad in ['8041-2237-01', '§1.4', '임원 보고용', '핵심 요약', '버전 v2.0', '3층 체제']:
        n = body2.count(bad)
        print(f'  {"OK(0)" if n == 0 else f"잔존 {n}!!"} — {bad}')
    h1 = next(p for p in d2.paragraphs if p.style.name == 'Heading 1' and p.text == '제안 개요')
    print('H1 크기/색:', h1.runs[0].font.size, h1.runs[0].font.color.rgb)


if __name__ == '__main__':
    main()
