# 기술설계서 v1.4 생성기 — v1.3.docx(진본)를 읽어 한국식 개조식 문서로 변환해 v1.4.docx로 저장.
# v1.3의 수기 편집을 보존하기 위해 재조판 없이 run 단위 텍스트만 치환한다(서식·표·볼드 유지).
# 변환 내용: ① 줄표(—) 제거(라벨 뒤 콜론, 문장 사이 마침표) ② 번역체·명사화 수정
#            ③ 합니다체 → 개조식(~함/~임/~됨) ④ 표지 버전 v1.4
# 실행: python3 docs/기술설계서_v14_생성.py  (레포 루트 기준)
import re
from docx import Document

SRC = 'docs/수수료정책플랫폼_기술설계서_v1.3.docx'
DST = 'docs/수수료정책플랫폼_기술설계서_v1.4.docx'

EXPLICIT = [
    ('버전 v1.3', '버전 v1.4'),
    ('하나의 개념이 가지는 속성일 뿐입니다', '한 개념의 속성일 뿐입니다'),
    ('구간표이면 가격 구간을 가집니다', '구간표이면 가격 구간이 있습니다'),
    ('상태를 가집니다', '상태를 거칩니다'),
    ('순으로 상세화합니다', '순으로 자세히 풀어 씁니다'),
    ('없어 탈락 — 협의가', '없어 탈락합니다. 협의가'),
    ('셀을 전개 — 계좌가 개설한 상품군만', '셀을 전개(계좌가 개설한 상품군만)'),
    ('일배치가 정합성의 진본.', '일배치가 정합성의 진본입니다.'),
    ('통합 관리하고, 고객이', '통합 관리하고 고객이'),
    ('저장하고, 매일 배치로 계좌별', '저장하고 매일 배치로 계좌별'),
    ('산출하고, 부담주체별로', '산출하고 부담주체별로'),
    ('8041-2237-01 의', '8041-2237-01의'),
    ('6015-8890-42 는', '6015-8890-42는'),
    ('— 이상 —', '이상.'),
]
GAEJOSIK_SPECIFIC = [
    ('않았습니다', '않았음'), ('있었습니다', '있었음'),
    ('않습니다', '않음'), ('있습니다', '있음'), ('없습니다', '없음'),
    ('다릅니다', '다름'), ('같습니다', '같음'), ('아닙니다', '아님'),
    ('거칩니다', '거침'), ('다룹니다', '다룸'), ('나뉩니다', '나뉨'),
    ('둡니다', '둠'), ('씁니다', '씀'), ('이긴다', '이김'),
]
GAEJOSIK_GENERIC = [
    (re.compile(r'([가-힣]+)깁니다'), r'\1김'),   # 남깁니다→남김, 이깁니다→이김
    (re.compile(r'([가-힣]+)킵니다'), r'\1킴'),   # 경쟁시킵니다→경쟁시킴
    (re.compile(r'([가-힣]*)합니다'), r'\1함'),
    (re.compile(r'([가-힣]+)됩니다'), r'\1됨'),
    (re.compile(r'([가-힣]+)입니다'), r'\1임'),   # 쓰입니다→쓰임 포함
]


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    yield p
    for s in doc.sections:
        for hf in (s.header, s.footer):
            for p in hf.paragraphs:
                yield p


def main():
    d = Document(SRC)

    for p in iter_paragraphs(d):                      # ① 명시 교체(줄표 일괄 치환보다 먼저)
        for r in p.runs:
            for old, new in EXPLICIT:
                if old in r.text:
                    r.text = r.text.replace(old, new)

    for p in iter_paragraphs(d):                      # ② 문장 경계 줄표
        for r in p.runs:
            r.text = r.text.replace('니다 — ', '니다. ')

    for p in iter_paragraphs(d):                      # ③ 라벨 줄표 — 단락 첫 번째는 콜론, 이후는 마침표
        first = False
        for r in p.runs:
            while ' — ' in r.text:
                if not first:
                    r.text = r.text.replace(' — ', ': ', 1)
                    first = True
                else:
                    r.text = r.text.replace(' — ', '. ', 1)

    for t in d.tables:                                # ④ 표 빈칸 마커 "—" → "-"
        for row in t.rows:
            for c in row.cells:
                if c.text.strip() == '—':
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.text = r.text.replace('—', '-')

    for p in iter_paragraphs(d):                      # ⑤ 합니다체 → 개조식
        for r in p.runs:
            s = r.text
            for a, b in GAEJOSIK_SPECIFIC:
                s = s.replace(a, b)
            for pat, rep in GAEJOSIK_GENERIC:
                s = pat.sub(rep, s)
            s = s.replace('”입니다', '”임').replace(')입니다', ')임')
            r.text = s

    d.save(DST)

    d2 = Document(DST)                                # 검증
    texts = [p.text for p in d2.paragraphs]
    for t in d2.tables:
        for row in t.rows:
            for c in row.cells:
                texts.append(c.text)
    bad = [x for x in texts if '니다' in x or '—' in x or '–' in x]
    print(f'저장: {DST}')
    print(f'검증 — 잔존 합니다체·줄표 {len(bad)}건 / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)}')
    if bad:
        for x in bad:
            print('  |', x.strip()[:80])


if __name__ == '__main__':
    main()
