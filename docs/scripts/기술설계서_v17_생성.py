# 기술설계서 v1.7 생성기 — v1.6.docx를 읽어 이번 설계 확정분을 additive 반영해 v1.7.docx로 저장.
#   ① 용어 정비 — 부여관계 → 편입내역(부여ID → 편입ID, 단독 '부여' → '편입'), 용어표에 항목 추가
#   ② §1.4 보강 — 구체 셀은 승자가 같아도 저장(자기완결 원칙)
#   ③ §10.1 ② 정밀화 — 셀 전개는 한정 축들의 곱집합(교차 셀 포함)
#   ④ §10.2 말미 — 증분·delta 완결성 불변식(승자 입력은 기록과 날짜뿐)
#   ⑤ §10.4 신설 — 수십억 판의 후보·우선순위 처리(승인 시 순위값 저장·후보 색인·집합 산출)
#   ⑥ 표지 버전 v1.7 · 판 이력 갱신
# 체인: v1.3(진본) → v14 → v15 → v16 → v17. 실행: python3 docs/scripts/기술설계서_v17_생성.py (레포 루트 기준)
from docx import Document

SRC = 'docs/수수료정책플랫폼_기술설계서_v1.6.docx'
DST = 'docs/수수료정책플랫폼_기술설계서_v1.7.docx'


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_everywhere(doc, old, new):
    """본문·표의 모든 단락에서 치환. 런 경계에 걸린 경우만 단락을 평탄화한다."""
    n = 0
    for p in all_paragraphs(doc):
        if old not in p.text:
            continue
        for r in p.runs:
            if old in r.text:
                n += r.text.count(old)
                r.text = r.text.replace(old, new)
        if old in p.text:  # 런 경계에 걸림 → 첫 런에 몰아넣기
            txt = p.text.replace(old, new)
            n += 1
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = txt
            else:
                p.add_run(txt)
    return n


def main():
    d = Document(SRC)

    # ⑥-1 표지 버전 (런이 '버전'+' v1.6'으로 갈라져 있어 단락 단위로 찾는다)
    for p in d.paragraphs:
        if '버전 v1.6' in p.text:
            for r in p.runs:
                if 'v1.6' in r.text:
                    r.text = r.text.replace('v1.6', 'v1.7')

    # ① 용어 정비 — '부여'(받침 없음) → '편입'(받침 있음)이라 조사도 함께 바꾼다. 긴 토큰부터.
    pairs = [
        ('부여관계로', '편입내역으로'), ('부여관계와', '편입내역과'),
        ('부여관계가', '편입내역이'), ('부여관계는', '편입내역은'), ('부여관계를', '편입내역을'),
        ('부여관계', '편입내역'), ('부여ID', '편입ID'),
        ('부여로', '편입으로'), ('부여가', '편입이'), ('부여는', '편입은'), ('부여를', '편입을'),
        ('부여', '편입'),
    ]
    counts = {old: replace_everywhere(d, old, new) for old, new in pairs}
    c1 = sum(v for k, v in counts.items() if k.startswith('부여관계'))
    c2 = counts['부여ID']
    c3 = sum(v for k, v in counts.items() if not k.startswith('부여관계') and k != '부여ID')

    # 목록(글머리) 스타일 견본 — 기존 a0 단락의 스타일을 빌린다
    bullet_style = next(p.style for p in d.paragraphs
                        if p.style.style_id == 'a0' and p.text.strip())

    # ② §1.4 보강 — "2. 데이터 모델" H1 앞(= §1.4 끝)에 문단 추가
    h2 = next(p for p in d.paragraphs
              if p.style.name == 'Heading 1' and p.text.startswith('2. 데이터 모델'))
    h2.insert_paragraph_before(
        '구체 셀은 승자가 같아도 저장함(자기완결 원칙). 세션·채널처럼 "전체" 칸과 구체 값 칸이 함께 있는 '
        '축에서, 구체 셀의 승자가 "전체" 셀과 같더라도 행을 생략하지 않음. 원장 조회 계약을 "구체 행이 '
        '있으면 그것이 정답"이라는 한 줄로 유지하기 위한 조치이며, 이런 중복 행은 채널·세션 한정 정책이 '
        '경쟁에서 진 셀에서만 생겨 규모 영향이 미미함.')

    # ③ §10.1 ② 셀 전개 문구 정밀화
    cell_p = next(p for p in d.paragraphs if p.text.startswith('계좌 × ('))
    for r in cell_p.runs:
        r.text = ''
    cell_p.runs[0].text = (
        '계좌 × (자산군·거래소·조회구분·품목[파생]) 셀을 전개(계좌가 개설한 상품군만). 세션·채널'
        '(주식형은 거래소 포함)은 "전체" 한 칸으로 압축하되, 해당 축을 한정하는 활성 정책이 있으면 '
        '한정 값들의 곱집합으로 셀을 쪼갬. 서로 다른 축을 한정한 정책 둘이 겹치는 교차 셀'
        '(예: 야간×MTS)도 이때 만들어져, 어느 체결이 와도 완결된 승자가 판에 실림')

    # ④ §10.2 말미 불변식 — "10.3" H2 앞에 문단 추가
    h103 = next(p for p in d.paragraphs if p.text.startswith('10.3 규모 확장'))
    h103.insert_paragraph_before(
        '증분·delta가 판의 모든 변화를 빠짐없이 잡는 근거는 다음 불변식임: 배정판 승자를 바꾸는 입력은 '
        '기록(정책 상태·편입내역·계좌 플래그)과 날짜뿐이며, 계좌 지표는 신청·승인 절차를 거쳐서만 판에 '
        '반영됨. 협의가 안전한 이유가 이것으로, 지표가 임계값을 넘어도 판은 그대로이고 승인이라는 기록이 '
        '생겨야 증분이 돎. 지표 충족 시 자동 편입하는 이벤트를 도입하면 이 불변식이 깨지므로, 그 경우 '
        '지표 적재 배치가 임계값을 넘나든 계좌를 뽑아 증분 계기에 추가해야 함.')

    # ⑤ §10.4 신설 — "11. 등록" H1 앞
    h11 = next(p for p in d.paragraphs
               if p.style.name == 'Heading 1' and p.text.startswith('11. 등록'))
    h11.insert_paragraph_before('10.4 수십억 판의 후보·우선순위 처리: 계산은 승인 때, 산출은 집합으로',
                                style='Heading 2')
    h11.insert_paragraph_before(
        '우선순위 규칙(§5.1)은 그대로 두고, 계산 시점과 산출 방식만 규모에 맞게 옮김. 활성 정책은 수백 건 '
        '수준이라 순위 계산 자체는 싸며, 비용은 계좌×셀마다 후보를 다시 훑는 데서 나옴. 따라서 후보와 '
        '순위는 정책이 바뀔 때 확정하고, 대량 산출에서는 고르기만 함.')
    for t in [
        '순위값은 승인 때 계산해 저장: 정렬키(요율 최저 → 계층 → 범위 구체성 → 정책ID)를 정책 승인 시 '
        '한 번 계산해 규칙에 저장함. 요율표는 승인 후 바뀌지 않으므로 배치마다 재계산하지 않음',
        '조회키 조합별 후보 색인: 조합(자산군·조회구분·거래소·품목)마다 순위값 오름차순 후보 목록을 소형 '
        '테이블로 두고, 정책 승인·종료 시 그 정책이 걸리는 조합만 갱신함. 색인 전체가 메가바이트 규모라 '
        '유지 비용이 없음',
        '대량 산출은 계좌 순회 대신 집합 연산: 자격 게이트를 기록 조인으로 바꿈. 기본·일괄형은 개설 계좌 '
        '전개, 신청·가입·협의는 편입내역 조인, 휴면복귀는 계좌 플래그 조인. 붙인 결과에서 계좌×셀마다 '
        '순위값 최저 한 건만 남기고 키 정렬 순서로 적재함(§10.3 적재 원칙과 동일)',
        '경로 두 벌의 정합: 대량 경로(집합 연산)와 증분·검증 경로(셀 단위 확정)가 병존하나, 둘 다 저장된 '
        '같은 순위값에서 최저를 고르므로 우선순위 판단은 승인 시점 한 곳에만 있음. 주기 전체 검증이 두 '
        '경로의 결과를 상시 대조함',
    ]:
        h11.insert_paragraph_before(t, style=bullet_style)
    h11.insert_paragraph_before(
        '이 구조에서 지배 비용은 결과 정렬·적재로 수렴함. 2억 행 정렬 적재 6분 실측의 연장선이며, '
        '계좌 레인지 병렬로 더 줄어듦.')

    # ①-2 용어표에 편입내역 추가 / ⑥-2 판 이력 갱신
    for t in d.tables:
        head = t.rows[0].cells[0].text.strip()
        if head == '용어':
            row = t.add_row()
            row.cells[0].text = '편입내역'
            row.cells[1].text = ('계좌가 정책 대상으로 편입된 기록(이벤트 신청·가입, 협의 승인·연장). '
                                 '상태(요청/활성/반려/만료)와 유효기간을 계좌 단위로 가짐(구 명칭: 부여관계)')
        for r in t.rows:
            for c in r.cells:
                if c.text.startswith('v1.0 기술설계서'):
                    c.text = ('v1.0 기술설계서 → v1.1·v1.2 우선순위 저장 설계 → v1.3 전체 통합·한글화 → '
                              'v1.4~v1.6 배치 전략·배정판 저장 범위 → v1.7(본 문서, 편입내역 용어 정비·'
                              '셀 전개 곱집합·수십억 후보·우선순위 처리)')

    d.save(DST)

    # 검증
    d2 = Document(DST)
    body = '\n'.join(p.text for p in all_paragraphs(d2))
    print(f'저장: {DST} / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)}')
    print(f'치환: 부여관계 {c1} · 부여ID {c2} · 부여 {c3}')
    print('표지:', next(p.text for p in d2.paragraphs if '버전 v' in p.text))
    print('잔존 부여(용어표 구 명칭 1건이어야 함):', body.count('부여'))
    print('편입내역 등장:', body.count('편입내역'))
    for key in ('구체 셀은 승자가 같아도', '곱집합으로 셀을 쪼갬', '불변식임', '10.4 수십억'):
        print(f'  {"OK" if key in body else "누락!!"} — {key}')


if __name__ == '__main__':
    main()
