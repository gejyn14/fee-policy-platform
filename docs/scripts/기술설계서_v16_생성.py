# 기술설계서 v1.6 생성기 — v1.5.docx를 읽어 "배정판 저장 범위(우대분만)" 원칙을 추가해 v1.6.docx로 저장.
# v1.5 내용 대부분 유지, 구현 중 확정한 결정만 additive 반영:
#   ① §1.4 신설 — 배정판은 우대(이벤트·협의) 승자만 저장, 기본은 미저장·미스 시 직접 적용
#   ② §6 fallback 문장 수정 — 미스→기본은 예외가 아니라 정상 경로
#   ③ 표지 버전 v1.6
# 체인: v1.3(진본) → v14 → v15 → v16. 실행: python3 docs/scripts/기술설계서_v16_생성.py (레포 루트 기준)
from docx import Document

SRC = 'docs/수수료정책플랫폼_기술설계서_v1.5.docx'
DST = 'docs/수수료정책플랫폼_기술설계서_v1.6.docx'


def add_plain(anchor, text, style='Normal'):
    return anchor.insert_paragraph_before(text, style=style)


def main():
    d = Document(SRC)

    # 표지 버전
    for p in d.paragraphs:
        for r in p.runs:
            if '버전 v1.5' in r.text:
                r.text = r.text.replace('버전 v1.5', '버전 v1.6')

    # ① §1.4 삽입 — "2. 데이터 모델" H1 앞
    h2 = next(p for p in d.paragraphs
              if p.style.name == 'Heading 1' and p.text.startswith('2. 데이터 모델'))
    add_plain(h2, '1.4 배정판 저장 범위: 우대분만 담는다', style='Heading 2')
    add_plain(h2, '배정판에는 기본수수료 승자를 저장하지 않음. 이벤트·협의가 기본보다 유리해 이긴 셀만 '
                  '행으로 남기고, 기본이 이기는 대다수 셀은 배정판에 행이 없음.')
    add_plain(h2, '원장이 체결 조회 시 배정판에서 행을 찾지 못하면(대다수 경우) 그것이 정상이며, '
                  '기본수수료를 직접 적용함. 기본 요율표는 자산군·조회구분으로 결정되고 계좌마다 다르지 않아, '
                  '계좌×셀로 복제 저장할 이유가 없음.')
    add_plain(h2, '효과: 배정판이 “기본 대비 우대분”만 담아 억 단위에서 수백만 규모로 축소되고, '
                  '일배치·증분 산출도 우대를 받은 계좌·셀만 대상으로 하므로 그만큼 가벼워짐. '
                  '원장 관점에서는 “배정판에 있으면 우대, 없으면 기본”이라는 단순한 이분 규칙이 됨.')

    # ② §6 fallback 문장 수정
    for p in d.paragraphs:
        if p.text.strip().startswith('배정판에 맞는 행이 없으면'):
            for r in p.runs:
                r.text = ''
            p.runs[0].text = ('배정판에 맞는 행이 없으면(대다수 셀) 기본수수료를 직접 적용함. '
                              '이는 예외가 아니라 정상 경로임(배정판은 우대분만 담음, §1.4 참조). '
                              '원장 수정 범위는 이 조회 한 곳뿐임.')
            break

    d.save(DST)

    d2 = Document(DST)
    print(f'저장: {DST} / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)}')
    print('표지:', next(p.text for p in d2.paragraphs if '버전 v' in p.text))
    i = next(i for i, p in enumerate(d2.paragraphs) if p.text.startswith('1.4 배정판'))
    for p in d2.paragraphs[i:i+4]:
        print('  ·', p.text[:70])


if __name__ == '__main__':
    main()
