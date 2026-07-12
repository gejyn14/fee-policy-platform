# 기술설계서 v2.0 생성기 — v1.9.docx를 읽어 네 가지를 반영해 v2.0.docx로 저장.
#   (1) 용어 변경: 조회구분 → 상품구분 (전 문서)
#   (2) §조회키의 논리 구조: 8축 반영(유동성·매수매도 행 추가 + 설명 정정)
#   (3) §배정판 예시: 축 추가 반영(유동성·매수매도 칼럼) + 기본 행 제거(배정판은 우대분만) + 메이커 예시 추가
#   (4) 문서 재구성: 11.3~11.5(규모 배치·후보 색인·1차전 물리화)를 중간(조회키 뒤)으로 이동해
#       새 5장 '대규모 산출'을 만들고 이후 장 번호·상호참조·목차 일괄 정정.
#   장 번호 재배치(구→신): 0~4 유지, [신규 5=대규모 산출], 구5→6 … 구12→13, 부록 유지.
#   이동 소절(구→신): 11.5 물리화→5.1, 11.4 후보처리→5.2, 11.3 규모전략→5.3.
# 체인: … v18 → v19 → v20. 실행: python3 docs/scripts/기술설계서_v20_생성.py
import re
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = 'docs/수수료정책플랫폼_기술설계서_v1.9.docx'
DST = 'docs/수수료정책플랫폼_기술설계서_v2.0.docx'
IMG_CAND = 'docs/diagrams/v20_후보선정_효율.png'

CHMAP = {0: 0, 1: 1, 2: 2, 3: 3, 4: 4, 5: 6, 6: 7, 7: 8, 8: 9, 9: 10, 10: 11, 11: 12, 12: 13}
# 이동 소절 상호참조 특례(구 → 신)
SPECIAL_REF = {'§11.5': '§5.1', '§11.4': '§5.2', '§11.3': '§5.3'}

NEW_TOC = [
    '0. 수수료 정책 플랫폼 개요',
    '1. 아키텍처: 정책은 저장, 배정판은 사전 산출',
    '2. 정책의 여정: 등록에서 원장 조회까지',
    '3. 수수료 우선순위(1차전 경쟁 규칙): 세 정책이 한 순위에서 경쟁',
    '4. 조회키: 체결을 수수료 정책에 잇는 기준',
    '5. 대규모 산출: 1차전 물리화·후보 색인·배치 전략',
    '6. 데이터 모델: 무엇을 어디에 담는가',
    '7. 수수료 계산: 요율표 평가',
    '8. 개별 체결의 수수료 결정: 원장의 배정판 조회',
    '9. 대상 편입: 고객이 정책 대상이 되는 방식',
    '10. 이벤트 기간: 신청 가능기간과 적용기간',
    '11. 협의수수료: 신청·자격·승인·연장',
    '12. 배정판 산출·무효화: 일배치 + 수시 증분',
    '13. 등록 → 검증 → 승인 → 활성',
    '부록. 배정판 점조회 성능 실측',
]


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def set_text(p, text):
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def find(doc, pred):
    return next(p for p in all_paragraphs(doc) if pred(p))


def remap_ref_token(tok):
    if tok in SPECIAL_REF:
        return SPECIAL_REF[tok]
    m = re.match(r'§(\d+)(\.\d+)?', tok)
    ch = int(m.group(1)); rest = m.group(2) or ''
    return '§' + str(CHMAP.get(ch, ch)) + rest


def insert_col(table, after_idx, header, value):
    """after_idx 칼럼을 복제해 그 뒤에 새 칼럼 삽입(서식 보존). 헤더/값 텍스트 지정."""
    grid = table._tbl.tblGrid
    gcols = grid.findall(qn('w:gridCol'))
    grid.insert(list(grid).index(gcols[after_idx]) + 1, copy.deepcopy(gcols[after_idx]))
    for ri, row in enumerate(table.rows):
        tcs = row._tr.findall(qn('w:tc'))
        newc = copy.deepcopy(tcs[after_idx])
        tcs[after_idx].addnext(newc)
        ts = newc.findall('.//' + qn('w:t'))
        for t in ts:
            t.text = ''
        txt = header if ri == 0 else value
        if ts:
            ts[0].text = txt
        else:
            newc.find('.//' + qn('w:p')).append(newc.makeelement(qn('w:r'), {}))


def main():
    d = Document(SRC)
    h1_style = next(p.style for p in d.paragraphs if p.style.name == 'Heading 1')
    normal_style = next(p.style for p in d.paragraphs
                        if p.style.name == 'Normal' and p.text.strip())
    bullet_style = next(p.style for p in d.paragraphs
                        if p.style.style_id == 'a0' and p.text.strip())

    # ── 0) 표지 버전
    for p in d.paragraphs:
        if '버전 v1.9' in p.text:
            for r in p.runs:
                if 'v1.9' in r.text:
                    r.text = r.text.replace('v1.9', 'v2.0')

    # ── 1) 용어 변경: 조회구분 → 상품구분 (전 문서)
    ren = 0
    for p in all_paragraphs(d):
        if '조회구분' in p.text:
            ren += p.text.count('조회구분')
            for r in p.runs:
                if '조회구분' in r.text:
                    r.text = r.text.replace('조회구분', '상품구분')
            if '조회구분' in p.text:
                set_text(p, p.text.replace('조회구분', '상품구분'))

    # ── 2) §조회키의 논리 구조: 설명 정정 + 표에 유동성·매수매도 행 추가
    s55 = find(d, lambda p: p.text.startswith('주식은 종목 단위로 우대를 걸지 않으므로'))
    set_text(s55,
        '주식은 종목 단위로 우대를 걸지 않으므로 품목을 “전체”로 두고 나머지 축(거래소·상품구분·세션·채널·'
        '유동성·매수매도)으로 수수료를 결정함. 조회키는 8개 축으로 이뤄지며, 상품 유형에 따라 실제로 쓰이는 '
        '축이 달라짐(주식은 품목이 전체, 파생은 유동성·품목이 핵심). 이것이 규모를 결정적으로 줄이는 핵심 조치임.')
    t_logic = next(t for t in d.tables
                   if t.rows[0].cells[0].text.strip() == '항목(칼럼)'
                   and any('메이커' not in r.cells[0].text for r in t.rows))
    for label, val, col3, col4 in [
        ('유동성', '메이커 / 테이커 / 전체', '해당 없음(전체)', '○'),
        ('매수·매도', '매수 / 매도', '○', '○'),
    ]:
        row = t_logic.add_row()
        row.cells[0].text = label
        row.cells[1].text = val
        row.cells[2].text = col3
        row.cells[3].text = col4

    # ── 3) §배정판 예시 표: 칼럼 추가 + 기본 행 제거 + 메이커 예시 + 표2(전부 기본) 제거
    bind = next(t for t in d.tables
                if t.rows[0].cells[0].text.strip() == '계좌번호' and len(t.rows) >= 6)
    # 3-a) 기본 행 제거(정책구분 == 기본)
    for row in list(bind.rows[1:]):
        if row.cells[-1].text.strip() == '기본':
            row._tr.getparent().remove(row._tr)
    # 3-b) 주문채널(6) 뒤에 유동성·매수매도 칼럼 추가(기존 행은 '전체')
    insert_col(bind, 6, '유동성', '전체')
    insert_col(bind, 7, '매수매도', '전체')
    # 3-c) 메이커 우대 이벤트 예시 행(유동성 축이 셀을 쪼갠 예)
    vals = ['8041-2237-01', '해외파생', 'CME', '선물', '주간', 'ES', '전체', '메이커', '전체',
            '2026-07-01', '2026-09-30', 'ES 메이커 이벤트 $0.60', '이벤트']
    nr = bind.add_row()
    for k, v in enumerate(vals):
        nr.cells[k].text = v
    # 3-d) 전부-기본 계좌 표(2행짜리) 제거
    base_only = next(t for t in d.tables
                     if t.rows[0].cells[0].text.strip() == '계좌번호' and len(t.rows) == 2)
    base_only._tbl.getparent().remove(base_only._tbl)
    # 3-e) 참고 문구 + 불릿 정정
    note = find(d, lambda p: p.text.startswith('참고: 협의 없는 일반 계좌'))
    set_text(note,
        '참고: 협의·이벤트가 없는 일반 계좌(예: 6015-8890-42)는 배정판에 행이 전혀 없음. 모든 체결이 '
        '기본수수료로 직접 해석되므로 저장할 우대분이 없음(§1.4).')
    opt = find(d, lambda p: p.text.startswith('ES 옵션: 협의 범위 밖'))
    set_text(opt,
        'ES 옵션: 협의 범위 밖이라 배정판에 행이 없음 → 체결 시 기본($2.50)을 직접 적용(미스=기본). '
        '같은 품목이라도 선물·옵션은 상품구분이 달라 별개')
    etf = find(d, lambda p: p.text.startswith('해외 ETF/주식'))
    set_text(etf,
        '해외 ETF는 이벤트(0.09%)라 배정판에 행이 있고, 해외주식(일반)은 기본이라 행이 없음(상품구분 분리 효과)')
    etf._p.addnext(copy.deepcopy(etf._p))  # 새 불릿 자리 확보용 복제 후 텍스트 교체
    maker_bullet = etf._p.getnext()
    from docx.text.paragraph import Paragraph
    mb = Paragraph(maker_bullet, etf._parent)
    set_text(mb,
        'ES 선물 메이커: 메이커 우대 이벤트($0.60)가 협의($0.80)보다 싸 메이커 셀만 이김. '
        '유동성 축이 셀을 쪼갠 예로, 테이커·전체는 협의가 유지됨')

    # ── 3-f) 적용범위(정책 스코프) 축 목록에 8축 반영: 상품구분·유동성·매수매도 추가
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.startswith('정책ID · 자산군 · 거래소 · 세션 · 채널 · 품목'):
                    c.text = ('정책ID · 자산군 · 상품구분 · 거래소 · 세션 · 채널 · 품목 · 유동성 · 매수매도 · '
                              '제외품목 (각 항목은 “전체” 또는 값 목록)')
    gian = find(d, lambda p: p.text.startswith('기안: 적용범위'))
    set_text(gian,
        '기안: 적용범위(자산군·상품구분·거래소·세션·채널·품목·유동성·매수매도)·요율표·기간·대상을 지정')

    # ── 4) 11.3~11.5 추출 → 중간 5장. 먼저 블록/앵커 포착(구 텍스트 기준)
    body = d.element.body
    sectPr = body.find(qn('w:sectPr'))
    para_of = {p._p: p for p in d.paragraphs}
    children = [c for c in body if c is not sectPr]

    def block_of(start_p):
        i = children.index(start_p._p)
        out = [children[i]]
        for c in children[i + 1:]:
            if c.tag == qn('w:p'):
                pp = para_of.get(c)
                if pp is not None and pp.style.name in ('Heading 1', 'Heading 2'):
                    break
            out.append(c)
        return out

    blk_phys = block_of(find(d, lambda p: p.text.startswith('11.5 1차전의 물리화')))   # →5.1
    blk_cand = block_of(find(d, lambda p: p.text.startswith('11.4 수십억 판의 후보')))  # →5.2
    blk_scale = block_of(find(d, lambda p: p.text.startswith('11.3 규모 확장 배치')))   # →5.3

    # 4장(조회키) 마지막 요소 = 앵커
    h4 = find(d, lambda p: p.style.name == 'Heading 1' and p.text.startswith('4. 조회키'))
    h4i = children.index(h4._p)
    nxt = next(children.index(p._p) for p in d.paragraphs
               if p.style.name == 'Heading 1' and children.index(p._p) > h4i)
    anchor = children[nxt - 1]

    # ── 5) 상호참조 §토큰 일괄 정정(단일 패스)
    for p in all_paragraphs(d):
        if '§' in p.text:
            new = re.sub(r'§\d+(?:\.\d+)?', lambda m: remap_ref_token(m.group(0)), p.text)
            if new != p.text:
                set_text(p, new)

    # ── 6) 헤딩 앞자리 정정(CHMAP) — 이동 소절은 뒤에서 5.x로 덮음
    for p in all_paragraphs(d):
        if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
            m = re.match(r'(\d+)', p.text)
            if m and int(m.group(1)) in CHMAP:
                set_text(p, str(CHMAP[int(m.group(1))]) + p.text[m.end(1):])

    # 이동 소절 제목을 5.x로 확정
    set_text(find(d, lambda p: '1차전의 물리화' in p.text and p.style.name == 'Heading 2'),
             '5.1 1차전의 물리화: 색인 키 8축 완전 키화, 계산은 물리 테이블로')
    set_text(find(d, lambda p: '후보·우선순위 처리' in p.text and p.style.name == 'Heading 2'),
             '5.2 수십억 판의 후보·우선순위 처리: 계산은 승인 때, 산출은 집합으로')
    set_text(find(d, lambda p: '규모 확장 배치 전략' in p.text and p.style.name == 'Heading 2'),
             '5.3 규모 확장 배치 전략: 억 단위 판')

    # ── 7) 새 5장 H1 + 도입부 생성(끝에 붙였다 이동)
    new_elems = []
    h5 = d.add_paragraph('5. 대규모 산출: 1차전 물리화·후보 색인·배치 전략', style=h1_style)
    new_elems.append(h5._p)
    intro = d.add_paragraph(
        '정책이 승인되면 1차전이 계좌와 무관하게 승자 순위를 확정하고 물리 테이블에 적재함. 이 장은 그 산출을 '
        '억 단위 규모에서 어떻게 값싸게 유지하는지를 다룸(물리화 → 후보·우선순위 처리 → 규모 배치 전략). '
        '세부 배치 운영(일배치·수시 증분·무효화)은 §12에서 이어짐.', style=normal_style)
    new_elems.append(intro._p)

    # ── 8) 블록 재배치: 4장 뒤에 [5장 H1, 도입부, 물리화, 후보처리, 규모전략]
    order = new_elems + blk_phys + blk_cand + blk_scale
    cur = anchor
    for e in order:
        cur.addnext(e)
        cur = e

    # ── 8-b) §5.2에 경쟁 상대 선정·전개 대비 효율 보강(+도식). 5.3 앞에 삽입.
    h53 = find(d, lambda p: p.style.name == 'Heading 2' and p.text.startswith('5.3 규모 확장'))
    for t in [
        '경쟁 상대는 계좌를 보지 않고 조합 단위로 고름. 정책이 승인되면 순위 색인 재구성이 계좌 무관으로 돌며, '
        '조합마다 던지는 질문은 하나임: 이 조합에서 어떤 정책이 이 정책과 함께 줄을 서야 하는가. 활성 정책 '
        '전부를 그 조합에 대해 범위 판정(자산군·상품구분·거래소·품목… 순서로 매칭)에 걸어, 범위가 닿는 정책만 '
        '후보로 남기고 순위값 오름차순으로 색인에 심음.',
        '한 정책은 자기가 한정한 값의 조합에만 닿음. 예로 대체거래소(NXT)만 겨냥한 새 이벤트가 들어오면 거래소 '
        '축이 {전체, KRX}에서 {전체, KRX, NXT}로 늘고, 갱신은 NXT 조합 한 곳에서만 일어남. KRX·전체 조합은 '
        '손대지 않고, 정책의 순위값은 규칙 컬럼에 한 번만 적음. “걸리는 조합만 갱신”이 이렇게 지켜짐.',
    ]:
        h53.insert_paragraph_before(t, style=normal_style)
    picp = h53.insert_paragraph_before('')
    picp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picp.add_run().add_picture(IMG_CAND, width=Cm(16.6))
    cap = h53.insert_paragraph_before('[그림] 1차전 경쟁 상대 선정과 전개 대비 효율(정책이 부른 값만 색인)')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for t in [
        '이 방식이 “모든 경우의 수 전개”보다 결정적으로 싼 이유는 칸을 정책이 부른 값으로만 만들기 때문임. '
        '축마다 값이 수십~수백 개라 전부 곱하면 조합이 폭발하지만 그 대부분은 어떤 정책도 겨냥하지 않는 빈 칸임. '
        '그래서 두 번 거름: 아무도 한정하지 않은 값은 전체(*) 한 칸으로 뭉치고(1차), 후보가 하나도 없는 칸은 '
        '아예 저장하지 않음(2차).',
        '실측 예(국내주식 활성 정책 9건)에서, 8축을 전부 곱하는 대신 정책이 부른 값만 남기면 조합은 21칸이고 '
        '그중 후보가 있는 12칸에 총 30행만 실림(국내주식은 품목·세션·유동성·매수매도가 전부 전체라 실질적으로 '
        '거래소×채널로 갈림). 색인 크기는 계좌 수나 이론상 경우의 수가 아니라 정책이 얼마나 잘게 나누느냐에만 '
        '비례하므로, 억 단위 계좌에서도 전체 색인이 수백 KB에 머묾. 조회는 8축 키로 WHERE 조건을 걸고 순위 '
        '순서로 한 줄만 읽으므로(LIMIT 1) 테이블이 억 행이어도 인덱스로 O(1)임. 저장을 조금 더 쓰는 대가로'
        '(넓게 걸치는 정책은 조합마다 중복 저장) 조회 O(1)을 사는 거래임.',
    ]:
        h53.insert_paragraph_before(t, style=normal_style)

    # ── 9) 목차 재작성
    toc_h1 = find(d, lambda p: p.text.strip() == '목차')
    toc_items, seen = [], False
    for p in d.paragraphs:
        if p._p is toc_h1._p:
            seen = True
            continue
        if seen:
            if p.style.name == 'Heading 1':
                break
            if p.text.strip():
                toc_items.append(p)
    toc_style = toc_items[0].style
    zero_h1 = find(d, lambda p: p.style.name == 'Heading 1' and p.text.startswith('0. 수수료'))
    for i, title in enumerate(NEW_TOC):
        if i < len(toc_items):
            set_text(toc_items[i], title)
        else:
            zero_h1.insert_paragraph_before(title, style=toc_style)

    # ── 10) 판 이력 행 제거 — 외부 공유 최종안에는 내부 변경 이력을 넣지 않음
    for t in d.tables:
        for r in list(t.rows):
            if r.cells[0].text.strip() == '판 이력':
                r._tr.getparent().remove(r._tr)

    d.save(DST)

    # ── 검증
    d2 = Document(DST)

    def allp2():
        return all_paragraphs(d2)
    body2 = '\n'.join(p.text for p in allp2())
    print(f'저장: {DST} / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)} / 이미지 {len(d2.inline_shapes)}')
    print('표지:', next(p.text for p in d2.paragraphs if '버전 v' in p.text))
    print(f'조회구분 치환: {ren}건 / 잔존 조회구분: {body2.count("조회구분")}')
    print('H1 순서:')
    for p in d2.paragraphs:
        if p.style.name == 'Heading 1':
            print('   ', p.text[:44])
    for key in ['5. 대규모 산출', '5.1 1차전의 물리화', '5.2 수십억', '5.3 규모 확장',
                '유동성', '매수·매도', 'ES 메이커 이벤트', '§5.1', '§12', '§13']:
        print(f'  {"OK" if key in body2 else "누락!!"} — {key}')
    print('잔존 §11.:', len(re.findall(r'§11\.', body2)), '/ §7 참조:', len(re.findall(r'§7(?!\d)', body2)))
    b = next(t for t in d2.tables if t.rows[0].cells[0].text.strip() == '계좌번호')
    print(f'배정판 표: {len(b.rows)}행 x {len(b.rows[0].cells)}칼럼 / 헤더: '
          + ' | '.join(c.text for c in b.rows[0].cells))
    print('정책구분들:', [r.cells[-1].text for r in b.rows[1:]])


if __name__ == '__main__':
    main()
