# -*- coding: utf-8 -*-
# 개선 제안서 v2.0 생성기 — v1.0을 읽어 전면부를 맥킨지 SCR 구조로 재구성해 저장.
#   구조 근거: 맥킨지·BCG·베인의 제안 문서 표준 —
#     · Executive Summary는 결론 먼저(권고를 첫 문장에 단정형으로), 굵은 주장 + 들여쓴 근거 불릿.
#       굵은 글씨만 읽어도 논지 전체가 통해야 함. Resolution이 분량의 60~70%.
#     · 본문 스토리는 상황(Situation) → 문제(Complication) → 해결(Resolution).
#   아마존 6-pager는 서술형 산문 전제라 개조식 한국 보고서와 상극이므로 배제.
#   본문 장(0~13)·상호참조·표·도식은 그대로 두고 전면부만 바꿈. 한국식 개조식·타이포는 v1.0 그대로.
# 실행: python3 docs/scripts/개선제안서_v2_생성.py (레포 루트 기준)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = 'docs/수수료정책플랫폼_개선제안서_v1.0.docx'
DST = 'docs/수수료정책플랫폼_개선제안서_v2.0.docx'

KFONT = '맑은 고딕'
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
NAVY2 = RGBColor(0x2E, 0x4E, 0x75)
BODY = RGBColor(0x26, 0x26, 0x2B)


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def set_text(p, text):
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def kfont(run, size=None, bold=None, color=None):
    run.font.name = KFONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), KFONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def h1_rule(p):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F3A5F')
    pbdr.append(bottom)
    ppr.append(pbdr)


def main():
    d = Document(SRC)
    bullet_style = next(p.style for p in d.paragraphs
                        if p.style.style_id == 'a0' and p.text.strip())

    # ── 1) 표지: 버전 상향
    date_p = next(p for p in d.paragraphs if p.text.strip() == '2026년 7월')
    set_text(date_p, '2026년 7월  ·  v2.0')

    # ── 2) 기존 전면부 제거: 제안 개요 H1의 도입문·개요 표·전환문
    h_over = next(p for p in d.paragraphs if p.text.strip() == '제안 개요')
    intro = next(p for p in d.paragraphs
                 if p.text.startswith('수수료 정책 플랫폼은 기본·이벤트·협의'))
    trans = next(p for p in d.paragraphs if p.text.startswith('아래 본문은 이 제안을'))
    intro._p.getparent().remove(intro._p)
    trans._p.getparent().remove(trans._p)
    otbl = next(t for t in d.tables if t.rows[0].cells[0].text.strip() == '추진 배경')
    otbl._tbl.getparent().remove(otbl._tbl)

    # ── 3) 제안 요약 (결론 먼저) — 주요 용어 H2 앞에 삽입
    set_text(h_over, '제안 요약')
    terms_h2 = next(p for p in d.paragraphs if p.text.strip() == '주요 용어')

    def para(text, style=None, size=10.5, bold=False, color=BODY,
             before=None, after=None, anchor=None):
        tgt = anchor if anchor is not None else terms_h2
        p = tgt.insert_paragraph_before(text, style=style) if style \
            else tgt.insert_paragraph_before(text)
        for r in p.runs:
            kfont(r, size, bold, color)
        if before is not None:
            p.paragraph_format.space_before = Pt(before)
        if after is not None:
            p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.42
        return p

    # 권고 (첫 문장이 곧 결론)
    para('권고: 현행 등급제 수수료 체계를 "수수료 정책 플랫폼"으로 전환할 것을 제안함. 기본·이벤트·협의를 '
         '하나의 정책으로 통합하고, 계좌별 수수료 배정판을 미리 산출해 원장은 체결 시 한 줄만 조회하게 함.',
         size=11, bold=True, color=NAVY, before=4, after=8)

    # 근거 1~3: 굵은 주장 + 들여쓴 근거 (굵은 글씨만 읽어도 논지가 통하도록)
    blocks = [
        ('근거 ① 현행 등급제로는 늘어나는 우대 수수료 요구를 담지 못함.',
         ['등급제(계좌 등급 → 등급·상품 요율)는 이벤트·협의·최저가 보장·기간 한정을 애초에 표현하지 못함',
          '주요사 수수료 이벤트 조사 결과, 정책은 "상품군·거래소·상품구분·채널(파생은 품목) + 기간"으로 '
          '정형화 가능함(§0.3). 새 틀이 실수요를 담을 수 있다는 근거임']),
        ('근거 ② 정책만 저장하고 배정판을 미리 산출하면 억 단위 규모에서도 성립함.',
         ['계좌×품목 전개 저장(수백억 건) 대신 정책 수백 건만 저장하고, 배정판은 우대분만 담아 수백만 '
          '규모에 머묾(§1.2)',
          '배정판 점조회는 2억 행에서 중앙값 0.3ms를 실측으로 확인함(부록). 규모의 관건인 산출 배치는 '
          '수시 증분 + 일 delta로 감당함(§5.3)',
          '우선순위 계산은 정책 승인 시점 1회로 모으고(1차전 물리화), 이후 모든 경로는 저장된 순위를 '
          '읽기만 함(§5.1)']),
        ('근거 ③ 원장 변경은 배정판 조회 한 곳뿐이라 이행 부담이 작음.',
         ['체결 시점에는 계산이 없고, 배정판에 행이 없으면 기본수수료를 직접 적용하는 이분 규칙임(§8)',
          '어느 계좌에 언제 어떤 정책이 적용됐는지 배정 이력으로 전부 추적되어 민원 대응이 가능함(§6.4)']),
    ]
    for claim, subs in blocks:
        para(claim, size=10.5, bold=True, color=BODY, before=6)
        for s in subs:
            para(s, style=bullet_style)

    para('결정 요청 사항: 본 설계안 확정과 플랫폼 구축 착수 승인.',
         size=10.5, bold=True, color=NAVY, before=8, after=10)

    # ── 4) 제안 배경: 상황 → 문제 → 해결
    bg_h1 = para('제안 배경: 상황 → 문제 → 해결', style='Heading 1')
    for r in bg_h1.runs:
        kfont(r, 16, True, NAVY)
    h1_rule(bg_h1)
    bg_h1.paragraph_format.space_before = Pt(18)
    bg_h1.paragraph_format.space_after = Pt(8)

    def h2(text):
        p = para(text, style='Heading 2')
        for r in p.runs:
            kfont(r, 13, True, NAVY2)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
        return p

    h2('상황: 수수료 우대는 다양해지는데, 그릇은 등급제 하나임')
    for s in ['마케팅·영업은 거래소·세션·채널·품목 단위의 한시 우대와 계좌별 협의 수수료를 상시 운영하고자 함',
              '현행 체계는 계좌 등급이 등급·상품 요율을 정하는 단일 구조로, 모든 우대가 이 틀을 우회해야 함']:
        para(s, style=bullet_style)

    h2('문제: 등급제로는 담지 못하고, 흔한 대안인 전개 저장은 성립하지 않음')
    for s in ['등급제는 이벤트·협의·최저가 보장·기간 한정을 표현하지 못함',
              '계좌마다 모든 종목의 수수료를 미리 펼쳐 저장하는 방식은 수백억 건이 되어 저장·산출 모두 불가능함',
              '체결 시점마다 우선순위를 계산하는 방식은 원장에 규칙·우선순위 판단을 심어 원장 부담과 장애 '
              '위험을 키움']:
        para(s, style=bullet_style)

    h2('해결: 정책 통합 + 배정판 사전 산출 + 승인 시점 1차전')
    for s in ['기본·이벤트·협의를 하나의 "정책"으로 통합함. 이벤트·협의는 별개 기능이 아니라 정책의 속성임(§0.1)',
              '정책만 저장하고, 매일 배치와 수시 증분으로 계좌별 수수료 배정판을 미리 산출함. 원장은 체결 시 '
              '배정판 한 줄만 조회함(§1·§2)',
              '우선순위 경쟁은 정책 승인 시점의 1차전에서 한 번 확정해 물리 테이블(순위값·8축 후보 색인)로 '
              '저장함. 배치는 저장된 순위를 읽기만 하므로 규모(계좌 수)와 무관함(§5)',
              '주식은 종목 축을 제거하고 파생만 품목별로 관리해 단일 테이블로 감당함(§6.5)']:
        para(s, style=bullet_style)
    para('아래 본문이 이 해결안의 상세 근거임. 개요 → 아키텍처 → 정책의 여정 → 우선순위 → 조회키 → '
         '대규모 산출 → 데이터 모델 → 운영 순으로 풀어 씀.', after=10)

    d.save(DST)

    # ── 검증
    d2 = Document(DST)
    body2 = '\n'.join(p.text for p in all_paragraphs(d2))
    print(f'저장: {DST} / 단락 {len(d2.paragraphs)} / 표 {len(d2.tables)} / 이미지 {len(d2.inline_shapes)}')
    order = [p.text.strip()[:24] for p in d2.paragraphs
             if p.style.name in ('Heading 1', 'Heading 2')][:8]
    print('전면부 헤딩 순서:', order)
    for k in ['권고: 현행 등급제', '근거 ①', '근거 ②', '근거 ③', '결정 요청 사항',
              '상황: 수수료 우대는', '문제: 등급제로는', '해결: 정책 통합', 'v2.0']:
        print(f'  {"OK" if k in body2 else "누락!!"} — {k}')
    print('개요 표 잔존(0이어야):', body2.count('추진 배경'))


if __name__ == '__main__':
    main()
