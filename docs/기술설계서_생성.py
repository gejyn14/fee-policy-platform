# -*- coding: utf-8 -*-
"""수수료 정책 플랫폼 — 기술 설계서(v0.8).
구현(도메인 모델·해석 알고리즘·우선순위 산정·대상 편입·협의 상태머신) 기준의 기술 문서.
0장(왜 수수료 정책 플랫폼인가)은 기존 문안을 그대로 보존한다. 맑은 고딕 / 검정 제목."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = '맑은 고딕'
MONO = 'D2Coding'  # 없으면 맑은 고딕으로 대체됨
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x55, 0x55, 0x55)
BLUE = RGBColor(0x1F, 0x49, 0x7D)
DIR = '/Users/yujin-an/dev/fees/docs/diagrams'

doc = Document()
st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def set_ko(r, font=FONT):
    r.font.name = font
    for tag in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
        r._element.rPr.rFonts.set(qn(tag), font)


def h(text, level=1):
    p = doc.add_heading(level=level); r = p.add_run(text); set_ko(r); r.font.color.rgb = BLACK; return p


def para(text, color=None, italic=False, size=10.5, after=6):
    p = doc.add_paragraph(); r = p.add_run(text); set_ko(r); r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def mono(text, after=6):
    """코드/식별자 블록 — 고정폭 느낌."""
    p = doc.add_paragraph(); r = p.add_run(text); set_ko(r, MONO); r.font.size = Pt(9.5); r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(2); return p


def b(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(3)
    if ' :: ' in text:
        head, rest = text.split(' :: ', 1)
        r = p.add_run(head); set_ko(r); r.bold = True
        r2 = p.add_run(' — ' + rest); set_ko(r2)
    else:
        r = p.add_run(text); set_ko(r)
    return p


def num(items):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(3)
        if ' :: ' in t:
            head, rest = t.split(' :: ', 1)
            r = p.add_run(head); set_ko(r); r.bold = True
            r2 = p.add_run(' — ' + rest); set_ko(r2)
        else:
            r = p.add_run(t); set_ko(r)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(t.rows[0].cells):
        c.text = ''; r = c.paragraphs[0].add_run(headers[i]); set_ko(r); r.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(cells):
            c.text = ''; r = c.paragraphs[0].add_run(str(row[i])); set_ko(r); r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def pic(name, w=16.6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f'{DIR}/{name}.png', width=Cm(w))
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)


# ============================================================ 표지
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('수수료 정책 플랫폼'); set_ko(r); r.bold = True; r.font.size = Pt(24)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('기술 설계서'); set_ko(r); r.font.size = Pt(16); r.font.color.rgb = GREY
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run('버전 0.8  ·  2026-07-05  ·  시스템/아키텍처 기술 상세'); set_ko(r); r.font.size = Pt(10.5); r.font.color.rgb = GREY
doc.add_paragraph()
para('이 문서는 수수료 정책 플랫폼의 기술 설계를 구현 기준으로 상세히 기술한다. 아키텍처(해석형 코어), '
     '데이터 모델과 각 테이블의 컬럼, 조회키(feeKey) 구성과 매칭, 수수료 계산, 수수료 우선순위를 언제·어떻게 '
     '정하는지(정책 우선순위 사전 산정), 개별 체결의 수수료 결정 흐름(해석), 이벤트 신청·액션 발생 시 어떻게 '
     '정책 대상이 되는지(대상 편입), 타겟추출형의 적용 방식, 협의수수료 신청·승인·연장 상태머신, 캐시·무효화까지 '
     '다룬다. 독자는 업무개발·아키텍처 담당을 상정하며, 모듈·함수·테이블 수준의 기술 용어를 사용한다.')

h('용어', level=2)
table(['용어', '정의'], [
    ['조회키(feeKey)', '체결을 수수료 정책에 매핑하는 5차원 키 — 상품군·거래소·세션·주문채널·품목. 주식은 품목이 null(붕괴), 파생만 품목 유지.'],
    ['요율표(FeeSchedule)', '수수료 구성요소(자사·유관기관·세금)의 방식(정률/정액/구간표)·부담주체·요율·최소수수료 집합. 룰·협의가 참조.'],
    ['규칙(FeeRule)', '수수료 정책 1건. type은 BASE(기본) 또는 EVENT(이벤트). 적용범위·기간·대상·요율표를 가짐.'],
    ['협의 grant(NegoException)', '계좌 단위로 부여되는 협의수수료. 상태(요청/활성/반려)·자격(충족/예외)를 가진 오버레이.'],
    ['해석(resolve)', '체결 시 (계좌, feeKey)에 유효한 정책 후보를 모아 최저 수수료를 선택하는 과정.'],
    ['정책 우선순위 인덱스', '계좌 무관 정책(기본+전체 대상 이벤트)의 정적 최저가 랭킹. 룰 발효·만료 때만 재산정.'],
    ['지배관계', '우대 요율표가 기준선을 전(全) 가격구간에서 하회하는지에 대한 등록 시 검증.'],
    ['역마진', '회사부담 합계가 자사 수취분을 초과하는지에 대한 경고.'],
    ['해석 캐시(RESOLVED_CACHE)', '해석 결과를 (계좌, feeKey)로 저장하는 read-through 캐시.'],
], widths=[4.0, 11.5])
doc.add_page_break()

# ============================================================ 0 (보존)
h('0. 설계 관점 — 왜 “수수료 정책 플랫폼”인가', level=1)

h('0.1 이벤트 플랫폼이 아니라 수수료 정책 플랫폼', level=2)
para('이 시스템은 “이벤트”를 관리하는 도구가 아니라, 수수료를 정책으로 규율하는 플랫폼입니다. '
     '이벤트수수료와 협의수수료는 결국 “수수료 정책”이라는 하나의 개념이 가지는 속성일 뿐입니다.')
b('기본수수료 :: 상시 적용되는 표준 정책(기준선)')
b('이벤트수수료 :: 한시적·조건적 정책(대상 편입 방식이 다를 뿐)')
b('협의수수료 :: 개별 고객 정책(자격 조건 충족 시)')
para('따라서 “이벤트 등록”도 크게 보면 하나의 수수료 정책을 등록·승인·적용하는 행위이며, 문서 전반은 '
     '이 관점(정책의 등록 → 대상 편입 → 체결 시 적용)에서 기술합니다.')

h('0.2 수수료 정책의 대상이 되는 방법 — “이벤트(액션) 발생”', level=2)
para('여기서 말하는 이벤트는 행사(promotion)가 아니라 “액션의 발생”입니다. 금융뿐 아니라 이커머스 '
     '플랫폼까지 두루 살펴본 결과, 고객이 정책 대상이 되는 경로는 몇 가지로 수렴합니다.')
table(['대상 편입 경로', '설명', '이 시스템의 유형'], [
    ['쿠폰 적용 / 신청', '고객이 쿠폰을 적용하거나 우대를 신청 → 대상이 됨', '신청형'],
    ['가입(유입)', '이벤트 가입·상품 가입 액션', '가입형'],
    ['상태 변화 시그널', '휴면 계정 복구·회원가입·멤버십 해지 등 액션 발생', '휴면복귀형(및 확장)'],
    ['선제적 타겟 주입', '특정 타겟을 미리 뽑아 강제로 정책 주입(드물지만 필요)', '타겟추출형'],
    ['조건 충족', '특정 자격 조건을 충족한 고객에게 부여', '조건 충족형(협의수수료)'],
], widths=[3.2, 8.3, 4])
para('선제적으로 타겟을 작성해 강제 주입하는 경우는 드물지만 가능해야 하므로 타겟추출형으로 지원하고, '
     '자격 조건을 기준으로 부여하는 형태는 협의수수료가 담당합니다.')

h('0.3 증권사 수수료 정책의 실제 형태 — 정형화 근거', level=2)
para('증권사들의 실제 수수료 이벤트를 조사해, 현업이 선택할 수 있는 속성으로 정형화가 가능함을 확인했습니다. '
     '대부분의 증권사 수수료 정책은 다음 범위 안에 있었습니다.')
b('거래소별 / 세션별 수수료 할인 (주식)')
b('특정 품목에 대한 수수료 할인 (파생)')
b('여기에 기간·적용기한이 결합')
para('가장 선구적인 토스증권조차 이보다 더 세분화된 이벤트는 두고 있지 않았습니다. 즉 “상품군·거래소·세션·'
     '채널(파생은 품목) + 기간(신청 가능기간·적용기간)”이라는 축으로 현업 선택지를 정형화하면 실제 수요를 '
     '충분히 담을 수 있다는 결론이며, 본 문서의 구조·시스템은 이 근거 위에 설계되었습니다.')
doc.add_page_break()

# ============================================================ 1 아키텍처
h('1. 아키텍처 — 해석형 코어', level=1)
para('핵심 결정은 “계좌×품목의 수수료를 미리 펼쳐 저장하지 않는다”입니다. 대신 정책(규칙·요율표·협의)만 저장하고, '
     '체결 순간에 그 계좌·조회키에 유효한 정책을 해석(resolve)해 최저 수수료를 즉석 산정한 뒤, 결과만 캐시합니다.')
pic('d5_arch')

h('1.1 기본 원리', level=2)
b('금액 미저장 :: 규칙·요율표·협의만 보관. 체결 시 resolve(계좌, feeKey)가 후보를 모아 최저가를 계산')
b('read-through 캐시 :: 해석 결과를 (계좌, feeKey)로 저장. 재조회는 저장된 답 한 건만 읽음')
b('증분 무효화 :: 규칙 승인/연장·지표 변경 시 영향받는 scope/계좌의 캐시만 무효화 → 다음 조회에서 lazy 재해석')
b('결정성 :: 오늘 날짜는 상수 TODAY로 고정, Date.now/난수 미사용 — 동일 입력이면 항상 동일 결과(테스트·재현)')

h('1.2 전개형 vs 해석형 (규모 근거)', level=2)
table(['방식', '저장', '조회', '한계/이점'], [
    ['전개형(계좌×품목 사전 전개)', '수천만 계좌 × 수천 품목 = 수백억 행', '단순 조회(빠름)', '저장량 폭증·재전개 비용 → 대규모 불가'],
    ['해석형(체결 시 해석 + 캐시)', '규칙·요율표·협의(예외)만', '캐시 적중=1건 / 미스=소수 후보 계산', '저장 선형·변경은 증분 무효화'],
], widths=[4.5, 4.5, 3.3, 3.2])
b('주식은 조회키에서 품목(종목)을 제거해 카디널리티가 가장 큰 축을 없앰. 파생만 품목을 키에 유지')

h('1.3 모듈 구성', level=2)
para('코드는 순수 도메인 로직(React 비의존), 중앙 상태(store), 화면(React)으로 분리됩니다.')
table(['레이어', '주요 모듈', '책임'], [
    ['도메인 로직', 'types · feeKey · calc · resolve · dominance · policyRank · negoExtension · qualify · cache · binding', '수수료 정책 타입·조회키·계산·해석·검증(순수 함수)'],
    ['상태 관리', 'store/useStore(zustand) · store/mock', '규칙·요율표·협의·가입이력·지표 보유, 액션(resolveFee·policyPriority·submitNegoRequest 등)'],
    ['마스터데이터', 'masterdata/instruments · derive', '종목·파생 품목 마스터'],
    ['화면(UI)', 'Wizard · Dashboard · Approvals · NegoRequest · NegoApproval · Negotiated · FeeTrace · PolicyPriority · AccountView · Instrument*', '정책 등록·승인·협의·해석 시연·조회'],
], widths=[3, 6.5, 6])

h('1.4 플랫폼 / 원장 경계', level=2)
table(['영역', '책임'], [
    ['플랫폼', '정책 등록·검증·승인, 적용범위·기간 관리, 체결 시 해석, 협의 부여/연장, 캐시·무효화, 정책 우선순위 인덱스 유지'],
    ['원장', '체결·잔고 반영, 계좌 지표 산정·전달, 휴면복귀 등 시그널 발생, 해석 결과에 따른 실제 수수료 부과'],
], widths=[3, 12.5])
doc.add_page_break()

# ============================================================ 2 데이터 모델
h('2. 데이터 모델 — 테이블과 컬럼', level=1)
para('정책은 “요율표 + 규칙 + (계좌 단위) 협의/이력”으로 저장됩니다. 계좌마다 저장되는 것은 협의 grant와 가입/신청 '
     '이력뿐이며, 계좌×품목 전량 전개 테이블은 두지 않습니다.')
pic('d3_data')

h('2.1 요율표 계열', level=2)
para('요율표는 구성요소(1:N)로 이루어지고, 구성요소가 구간표(rate_type=구간표)이면 구간(1:N)을 가집니다.')
table(['테이블', '주요 컬럼', '설명'], [
    ['FEE_SCHEDULE', 'schedule_id(PK) · name', '요율표. 규칙·협의가 참조'],
    ['FEE_COMPONENT', 'schedule_id(FK) · seq · name · kind[자사/유관기관/세금] · payer[고객부과/회사부담/면제] · rate_type[정률/정액/구간표] · rate_bp · flat_amount · min_fee', '구성요소(요율 방식·부담주체·요율·최소수수료)'],
    ['FEE_RATE_BAND', 'schedule_id·seq(FK) · from · to · rate_bp · flat', '구간표의 가격구간. 한 구간에 정률+정액 동시 가능(예: 0.14%+13원)'],
], widths=[3.6, 10, 2])

h('2.2 규칙 계열', level=2)
para('규칙(FEE_RULE)의 type은 BASE 또는 EVENT입니다(협의수수료는 별도 규칙 타입이 아니라 계좌 grant로 관리 — 2.3). '
     'apply_mode는 대상 편입 방식(7장), benefit_*는 적용기간 방식(8장)입니다.')
table(['테이블', '주요 컬럼'], [
    ['FEE_RULE', 'rule_id(PK) · name · type[BASE/EVENT] · status[초안/승인대기/활성/반려/만료] · apply_mode[타겟추출형/신청형/가입형/휴면복귀형] · start_date · end_date · benefit_kind[캘린더/상대] · benefit_months · schedule_id(FK) · target_account_ids(옵션·타겟추출형 지정 대상)'],
    ['FEE_RULE_SCOPE', 'rule_id(FK) · asset_class · exchanges · sessions · channels · products · exclude_products  (각 축은 * 또는 값 목록)'],
], widths=[3.6, 12])

h('2.3 계좌 계열', level=2)
para('계좌 단위로 저장되는 유일한 정책 데이터가 협의 grant입니다. 이벤트 참여 이력은 ENROLLMENT, 계좌 지표는 '
     'ACCOUNT_METRIC에 담깁니다.')
table(['테이블', '주요 컬럼', '설명'], [
    ['NEGO_GRANT', 'account_id · scope(asset_class·exchanges·products…) · schedule_id · valid_from · valid_to · status[요청/활성/반려] · qualify[충족/예외] · reason · request_id · requested_by · requested_at · approved_at', '협의수수료 부여(계좌 단위). 상태머신은 9장'],
    ['ENROLLMENT', 'account_id · rule_id · enrolled_at · channel', '신청형·가입형 편입 이력, 상대형 혜택의 가입일 기준'],
    ['ACCOUNT_METRIC', 'account_id · avg_asset_6m · volume_6m · dormant_returned · grade', '자격 판정·휴면복귀 시그널. 원장이 정기 배치로 최근 6개월을 재산정·적재(→ 9.5)'],
    ['QUALIFY_POLICY', 'asset_class · metric[6개월평균자산/약정액] · threshold', '상품군별 협의 자격 기준(참조 데이터)'],
], widths=[3.2, 10.3, 2])

h('2.4 해석 산출물', level=2)
table(['테이블/구조', '주요 컬럼', '설명'], [
    ['RESOLVED_CACHE', 'account_id · fee_key · schedule_id · source_rule_id · source[협의/이벤트/기본] · computed_at', '해석 결과 캐시(계좌, feeKey). read-through·증분 무효화'],
    ['정책 우선순위 인덱스', 'fee_key 매칭 → [정책, rank 오름차순]', '계좌 무관 정책의 정적 최저가 랭킹. 룰 변경 때만 재산정(5장)'],
], widths=[3.6, 8, 4])

h('2.5 조회키(FEE_KEY) 논리 구조', level=2)
table(['컬럼', '값', '주식', '파생'], [
    ['asset_class', '상품군(국내주식/해외주식/국내파생/해외파생/금현물)', '○', '○'],
    ['exchange', '거래소(KRX/NXT/CME/NASDAQ…)', '○', '○'],
    ['session', '세션(프리/정규/애프터)', '○', '○'],
    ['channel', '주문채널(HTS/MTS/API/ARS/센터/반대매매)', '○', '○'],
    ['product', '품목(계약/종목 코드)', '✗ (null)', '○'],
], widths=[3, 8, 2.3, 2.3])
para('주식은 종목 단위로 우대를 걸지 않으므로 품목을 null로 붕괴시켜 (거래소·세션·채널)만으로 해석합니다. '
     '이는 카디널리티가 가장 큰 축을 제거해 캐시·인덱스 규모를 결정적으로 줄이는 조치입니다.')
doc.add_page_break()

# ============================================================ 3 조회키
h('3. 조회키(feeKey) 구성과 scope 매칭', level=1)

h('3.1 구성 — 체결에서 조회키로', level=2)
para('체결 이벤트가 들어오면 조회키를 조립합니다. 종목(상품)에서 상품군·거래소·(파생)품목코드를, 주문에서 세션·'
     '주문채널을 얻습니다. 계좌·체결가·수량은 feeKey에 들어가지 않습니다(계좌는 협의/이력 조회, 체결가는 금액 계산에 사용).')
mono('deriveFeeKey(product, session, channel) → buildFeeKey(assetClass, exchange, session, channel, isDerivative ? productCode : null)')
b('feeKeyString = asset_class|exchange|session|channel[|product]  — 캐시 완전일치 키')

h('3.2 scope 매칭 — 정책은 완전일치 해시가 아니다', level=2)
para('정책 후보 탐색은 feeKey를 정책의 적용범위(scope)에 차원별로 맞춰봅니다. 각 축은 * (전체) 또는 값 목록입니다.')
table(['차원', '매칭 규칙(scopeMatchesKey)'], [
    ['asset_class', '정확히 일치해야 함'],
    ['exchanges', '* 이거나, 목록이 feeKey.exchange를 포함'],
    ['sessions', '* 이거나, 목록이 feeKey.session을 포함'],
    ['channels', '* 이거나, 목록이 feeKey.channel을 포함'],
    ['products (파생만)', '* 이거나 목록이 feeKey.product를 포함, 그리고 exclude_products에 없을 것'],
], widths=[3.5, 12])
para('즉 하나의 정책은 여러 feeKey에 걸릴 수 있고(와일드카드), 하나의 feeKey에는 여러 정책이 매칭될 수 있습니다. '
     '반면 해석 캐시는 feeKeyString 완전일치 해시로 조회합니다(같은 계좌·같은 조회키의 재체결은 즉시 적중).')
doc.add_page_break()

# ============================================================ 4 계산
h('4. 수수료 계산 (요율표 평가)', level=1)
para('선택된 요율표에 체결(체결가·수량)을 대입해 구성요소별 금액을 산출하고, 부담주체별로 합산합니다.')

h('4.1 구성요소 방식', level=2)
table(['rate_type', '계산식(구성요소 금액)', '예'], [
    ['정률', '거래대금(notional) × rate_bp / 10000', '해외주식 8bp'],
    ['정액', 'flat_amount × 수량(qty)', '해외파생 계약당 $2.5'],
    ['구간표', '체결가가 속한 구간에서 (거래대금×rate_bp/10000) + (flat×수량)', 'KOSPI200옵션 0.14% + 13원'],
], widths=[2.5, 9, 4])
para('구간표는 한 구간에서 정률과 정액 add-on을 동시에 더합니다. 예컨대 국내파생 옵션은 상위 구간에서 요율이 '
     '0.147%로 낮아지지만 +78원 정액이 그 하락을 보정해 가격 대비 단조 증가를 유지합니다(교차 없음 — 5.3 참조). '
     '구성요소에 최소수수료(min_fee)가 있으면 계산값이 그 값보다 낮을 때 최소액으로 올립니다.')

h('4.2 부담주체·통화', level=2)
b('부담주체 :: 고객부과(고객이 냄) · 회사부담(회사가 냄, 고객총액 제외) · 면제(0)')
b('고객부과 합계 = payer=고객부과 구성요소 합. 회사부담 합계는 역마진 경고에 사용')
b('통화 :: 해외주식·해외파생 = USD($), 국내주식·국내파생·금현물 = KRW(원). 표시는 상품군 기준으로 결정')

h('4.3 표본 체결(probe grid)', level=2)
para('요율표끼리의 우열은 단일 가격이 아니라 “구간 경계 + 표본가”로 이루어진 공동 검사구간에서 비교합니다. '
     '이는 구간표가 섞였을 때 특정 가격만 보고 잘못된 승자를 고르는 것을 막습니다(지배관계 검증·해석 최저가에 공통 사용).')
doc.add_page_break()

# ============================================================ 5 우선순위
h('5. 수수료 우선순위 — 언제, 어떻게 정하는가', level=1)

h('5.1 계층 우선순위 — 협의 무조건 우선', level=2)
para('한 체결에는 협의·이벤트·기본이 동시에 걸릴 수 있습니다. 계층 우선순위는 다음과 같습니다.')
num([
    '협의(NEGOTIATED) :: 활성 협의 grant가 있으면 무조건 우선(요율/수수료액이 항상 더 낮음)',
    '이벤트(EVENT) :: 협의가 없을 때, 계좌 무관 정책 우선순위 인덱스에서 최저가',
    '기본(BASE) :: 항상 후보에 포함되는 기준선',
])
para('구현상 해석 정렬의 최상위 기준은 “협의 우선(negoFirst)”이며, 그 아래에서 이벤트·기본을 최저가로 비교합니다.')

h('5.2 계좌 무관 정책 우선순위 — 사전 산정', level=2)
para('“언제 우선순위를 정하는가”의 답: 계좌 무관 층(기본 + 전체 대상 이벤트)의 최저가 순위는 체결 때가 아니라 '
     '룰이 발효·만료될 때 미리 산정합니다. 이벤트가 소수이므로 feeKey별 작은 정적 랭킹으로 유지할 수 있습니다.')
b('대상 :: BASE 전부 + 타겟추출형(EVENT) 중 지정 계좌 없음(전체). 신청/가입/휴면복귀/지정계좌/상대형은 계좌 특정이라 제외')
b('rank :: 기준 체결(가격 100·수량 10)의 고객부과 총액. 교차가 없으므로 이 한 점 순위가 전 가격에서 성립')
b('winnerFor(feeKey) :: 적용범위가 그 feeKey에 맞는 계좌 무관 정책 중 rank 최소를 즉시 룩업(재계산 없음)')
b('재산정 시점 :: 룰 발효/만료(활성 집합 변화) 시에만. 체결마다 다시 계산하지 않음')
para('체결 시에는 이 인덱스에서 계좌 무관 최저가를 룩업하고, 계좌에 협의가 있으면 그 값만 얹어 확정합니다. '
     '협의가 없는 대다수 계좌는 인덱스 룩업 한 번으로 끝납니다.')

h('5.3 승자가 가격에 따라 바뀌지 않는 근거', level=2)
para('같은 feeKey 안에서 후보들의 최저가 승자는 체결가와 무관합니다. 근거는 다음과 같습니다.')
b('구조 불혼합 :: 파생은 정액(계약당), 주식·국내파생은 정률. 같은 feeKey에 정률/정액이 섞이지 않음')
b('구간표 단조 :: 옵션 구간표의 정액 add-on이 요율 하락을 보정해 가격 대비 단조 증가(경계 dip 없음)')
b('지배관계 강제 :: 우대는 기준선을 전 구간에서 하회해야 성립 → 기본 vs 이벤트도 교차하지 않음')
b('품목별 feeKey :: 파생은 품목마다 feeKey가 달라 서로 겹치지 않음')
para('유일한 예외는 같은 feeKey에 서로 지배하지 않는 할인 2개가 동시에 걸리는 경우인데(가격대별로 유불리가 갈림), '
     '이벤트가 소수라 드물고 “동일 feeKey 중복 할인 금지” 정책 또는 후보 집합 캐싱으로 처리합니다.')

h('5.4 계층 내 동률 tie-break', level=2)
para('같은 계층에서 고객 부담이 동일한 후보가 여럿이면, 언제 조회해도 같은 결과가 나오도록 아래 순서로 확정합니다.')
num([
    '더 구체적인 적용범위 우선 :: 세션·채널을 한정한 이벤트가 전체 이벤트보다 우선',
    '더 최근에 시작된 규칙',
    '규칙 식별자 순서(결정성 보장)',
])
doc.add_page_break()

# ============================================================ 6 해석 흐름
h('6. 개별 체결의 수수료 결정 흐름 (resolve)', level=1)
para('개별 유저의 체결에 대해 수수료가 정해지는 과정을 단계로 봅니다. 각 단계에서 어떤 테이블을 어떤 키로 읽는지가 핵심입니다.')
pic('d6_flow', w=13.5)

h('6.1 단계와 조회 테이블', level=2)
table(['단계', '읽는 테이블 / 인덱스', '키', '반환'], [
    ['① 조회키 구성', '(조회 없음)', '체결(상품·세션·채널)', 'feeKey'],
    ['② 협의 예외(최우선)', 'NEGO_GRANT', 'account_id + scope 매칭·활성', '협의 요율표 또는 없음'],
    ['③ 계좌 무관 최저가', '정책 우선순위 인덱스', 'feeKey', '이벤트/기본 최저가 정책'],
    ['④ 가입/신청 이력', 'ENROLLMENT', 'account_id', '신청·가입 이벤트 대상 여부, 가입일'],
    ['⑤ 최저가 확정', '(②③④ 비교)', '—', '최종 출처·요율표'],
    ['⑥ 요율표 평가', 'FEE_SCHEDULE·COMPONENT·RATE_BAND', 'schedule_id', '자사·유관기관·세금 금액'],
    ['⑦ 원장 반영', '(원장)', '—', '실제 부과·잔고 반영'],
], widths=[3, 5.5, 4, 3])
para('②에서 활성 협의가 있으면 그것이 승자입니다(협의 무조건 우선). 없으면 ③의 계좌 무관 최저가와 ④로 활성화된 '
     '계좌 특정 이벤트를 포함해 ⑤에서 최저가를 확정합니다. ⑥은 미리 저장하지 않고 요율표×체결로 즉석 계산합니다.')
pic('d9_swimlane', w=16.8)

h('6.2 후보 수집과 게이트', level=2)
para('resolve는 세 종류의 후보를 모읍니다.')
b('협의 :: NEGO_GRANT 중 account_id 일치 · status=활성 · valid_from~valid_to 내 · scope가 feeKey에 매칭')
b('이벤트 :: 정책 인덱스에서 feeKey에 매칭되는 규칙 중 대상 게이트(isTarget)·기간 게이트(isBenefitActive) 통과분')
b('기본 :: feeKey에 매칭되는 BASE 요율표(항상 후보)')
para('대상 게이트: 타겟추출형=전 계좌(또는 지정 계좌) / 신청·가입=ENROLLMENT 이력 존재 / 휴면복귀=복귀 계좌. '
     '기간 게이트: 캘린더형=규칙 기간 내 / 상대형=가입일~가입일+N(신청 마감과 무관).')

h('6.3 최저가 비교와 캐시', level=2)
b('정렬 :: ① 협의 우선(negoFirst) → ② 공동 probe grid 평균 고객부담 최소 → ③ 계층(이벤트>기본) → ④ 구체성 → ⑤ 시작일 → ⑥ id')
b('캐시 :: 결과를 (계좌, feeKey)로 저장. 첫 조회는 미스(계산 후 저장), 이후는 적중(저장된 답 사용)')
b('무효화 :: 규칙 승인/연장·지표 변경 등으로 답이 달라질 수 있으면 영향 범위(scope 또는 계좌)만 무효화')
pic('d8_resolve', w=13.5)
doc.add_page_break()

# ============================================================ 7 대상 편입
h('7. 대상 편입 — 이벤트 신청/액션 발생 시 정책 대상이 되는 방식', level=1)
para('이벤트수수료는 “누가 대상인가”가 유형마다 다릅니다. 대상 여부는 체결 시 게이트(isTarget)로 판정되며, 그 근거 '
     '데이터(ENROLLMENT·지표·지정 목록)가 어디에 쌓이는지가 유형을 구분합니다.')
pic('d4_apply')
table(['유형(apply_mode)', '대상이 되는 조건', '근거 데이터', '고객/시스템 행위'], [
    ['타겟추출형', '적용범위(또는 지정 계좌)에 맞으면 자동', '규칙 scope / target_account_ids', '없음(또는 사전 타겟 주입)'],
    ['신청형', '신청→승인된 계좌만', 'ENROLLMENT(신청 채널·시점)', '고객 신청 → 심사 승인'],
    ['가입형', '이벤트에 가입(유입)한 계좌만', 'ENROLLMENT(가입일)', '고객 가입 액션'],
    ['휴면복귀형', '휴면에서 복귀한 계좌 자동', 'ACCOUNT_METRIC.dormant_returned', '없음(원장 시그널 감지)'],
    ['조건 충족(협의)', '자격 조건 충족 시 부여', 'NEGO_GRANT + QUALIFY_POLICY', '고객/PB 신청 → 자격 판정 → 승인'],
], widths=[2.8, 4.5, 4.2, 4])

h('7.1 신청형 — 신청 이력 기반', level=2)
para('고객이 채널(HTS/MTS/센터 등)로 신청하고 승인되면 ENROLLMENT에 (계좌, 규칙, 신청 채널·시점)이 남습니다. '
     '체결 시 대상 게이트는 이 이력의 존재를 확인합니다. 채널을 한정한 이벤트는 신청 채널이 판정에 사용됩니다.')

h('7.2 가입형 — 가입일이 혜택 시작점', level=2)
para('이벤트에 가입(유입)하면 ENROLLMENT에 가입일이 남습니다. 적용기간이 상대형이면 이 가입일이 혜택 시작 기준이 '
     '되어 고객마다 종료일이 달라집니다(8장).')

h('7.3 휴면복귀형 — 시그널 즉시', level=2)
para('원장이 휴면 계좌의 복귀를 감지하면 계좌 지표(dormant_returned)가 서고, 배치 없이 즉시 대상이 됩니다.')

h('7.4 타겟추출형 — 어떻게 적용되는가', level=2)
para('타겟추출형은 “우리가 임의로 타겟을 뽑아 특정 정책을 적용”하는 형태로, 두 가지 대상 범위를 지원합니다.')
b('전체(지정 계좌 없음) :: 적용범위(scope)에 맞는 모든 계좌가 자동 대상. 이 유형만이 계좌 무관 정책 우선순위 '
  '인덱스에 포함됩니다(5.2) — 체결 시 scope 매칭만으로 최저가 후보가 됩니다')
b('지정 계좌(target_account_ids) :: CSV/Excel로 계좌 리스트를 업로드해 특정 타겟에만 주입. 대상 게이트가 '
  '계좌 소속을 확인(계좌 특정이므로 계좌 무관 인덱스에서는 제외)')
para('즉 타겟추출형은 “계좌 행위 없이 시스템이 대상을 정하는” 경로입니다. 전체 대상은 정형화된 상시 우대처럼 '
     '인덱스로 사전 산정되고, 지정 대상은 업로드한 목록에 대한 게이트로 적용됩니다.')

h('7.5 조건 충족형(협의) — 자격 기반', level=2)
para('협의는 QUALIFY_POLICY(상품군별 지표·임계값) 기준으로 자격을 자동 판정하고, 충족(또는 영업 예외)하면 '
     'NEGO_GRANT로 부여합니다(9장). 대상 편입이 “행위”가 아니라 “자격”인 유일한 경로입니다.')
doc.add_page_break()

# ============================================================ 8 기간
h('8. 이벤트 기간 — 신청 가능기간과 적용기간', level=1)
para('이벤트에는 서로 다른 두 기간이 있습니다. 혼동을 막기 위해 분리해 저장합니다.')
b('신청/유입 가능기간 :: 고객이 이벤트에 들어올 수 있는 캘린더 구간(start_date~end_date)')
b('적용기간(혜택 기간) :: 실제 우대가 적용되는 구간. 규칙마다 두 방식 중 선택(benefit_kind)')
pic('d2_period')
table(['benefit_kind', '적용기간 산정', '특징'], [
    ['캘린더', '규칙 종료일(end_date)까지', '모든 고객이 같은 날 종료'],
    ['상대(+N개월)', '가입일 ~ 가입일 + benefit_months', '고객마다 종료일 상이. 신청 마감 지나도 잔여기간 유지'],
], widths=[3, 8.5, 4])
para('상대형은 계좌마다 가입일이 다르므로 종료일이 제각각입니다. 체결 시 기간 게이트(isBenefitActive)는 상대형이면 '
     'ENROLLMENT의 가입일을 읽어 가입일+N을 계산하고, 그 범위 안일 때만 후보에 포함합니다.')
doc.add_page_break()

# ============================================================ 9 협의
h('9. 협의수수료 — 신청 · 자격 · 승인 · 연장', level=1)
para('협의수수료는 이벤트 규칙과 분리되어, 계좌 단위 grant(NEGO_GRANT)로 관리됩니다. grant는 상태·자격을 갖는 '
     '상태머신입니다.')
pic('d10_nego_flow', w=16.8)

h('9.1 상태머신', level=2)
table(['status', '의미', '전이'], [
    ['요청', '신청 접수(자격 판정 결과 포함)', '승인 → 활성 / 반려 → 반려'],
    ['활성', '유효기간 내 부여됨(해석 후보)', '연장 → valid_to 연장 / 탈락 → 반려'],
    ['반려', '거절 또는 탈락', '(종료)'],
], widths=[2.5, 7, 6])
b('qualify :: 충족(자격 통과) 또는 예외(미충족이나 영업상 우대 필요로 bypass)')

h('9.2 신청과 자격 자동판정', level=2)
para('신청 화면에서 상품군·적용범위·요율표(선택·수정)와 계좌 리스트를 입력하면, 계좌별로 상품군 표준 자격을 '
     '자동 판정합니다.')
b('자격 판정 :: QUALIFY_POLICY의 지표(6개월 평균자산/약정액)와 임계값을 계좌 지표와 비교(qualifyOf)')
b('영업 예외(bypass) :: 미충족이라도 영업상 필요한 계좌는 사유를 달아 예외로 신청 가능(qualify=예외)')
b('요율표 수정 :: 선택한 표준 요율표를 상속하되, 자사 요율을 정률(bp)/정액(원·$) 단위로 조정 가능. 구간표는 선택만')
b('요청 생성 :: [협수 요청] 시 계좌별 grant를 status=요청으로 생성(request_id로 묶음)')

h('9.3 승인', level=2)
para('별도 승인 탭에서 요청을 request_id 단위로 확인하고 승인/반려합니다. 승인 시 status=활성, valid_from=승인일, '
     'valid_to=승인일+1년으로 설정되고, 해당 (계좌, feeKey) 캐시가 무효화됩니다.')

h('9.4 연장 대상 산출', level=2)
para('만료가 다가오면 계좌를 일일이 보지 않고 그룹 단위로 연장 대상을 산출합니다(classifyNegoExtension). '
     '정기 배치가 아니라 만료 임박 시 자동 산출 → 담당자 확인 → 일괄 승인하는 요청형입니다.')
b('그룹 축 :: 주식형 협의는 상품군(국내주식·해외주식·금현물), 파생 협의는 품목 단위')
b('재평가 :: 활성 grant를 상품군 자격 기준으로 다시 평가 → 유지(자격 충족 또는 예외) / 탈락(더는 미충족)')
b('적용 :: 유지 → valid_to 연장, 탈락 → status=반려(해지 대상). 탈락 대상이 함께 노출됨')
pic('d7_nego')

h('9.5 6개월 실적 지표 — 언제 계산되는가', level=2)
para('협의 자격 판정에 쓰이는 지난 6개월 실적(평균자산·약정액)은 신청·연장 순간에 즉석 계산하지 않습니다. 원장이 '
     '정기 배치로 최근 6개월(rolling window)을 재산정해 ACCOUNT_METRIC에 미리 적재하고, 플랫폼은 그 최신 '
     '스냅샷을 참조만 합니다. 즉 “실적 계산”은 배치에서 한 번 이루어지고, 신청·연장 판정은 그 결과를 읽습니다.')
table(['시점', '무엇을', '실적 계산 여부'], [
    ['정기 배치(원장)', '최근 6개월 평균자산·약정액 재산정 → ACCOUNT_METRIC 적재', '여기서만 계산(예: 일 또는 월 주기)'],
    ['협의 신청 시', 'qualifyOf가 계좌의 최신 지표 스냅샷을 상품군 임계값과 비교 → 충족/미충족', '재계산 없음(스냅샷 조회)'],
    ['연장 대상 산출 시', 'classifyNegoExtension이 최신 스냅샷 기준으로 유지/탈락 판정', '재계산 없음(스냅샷 조회)'],
], widths=[3.2, 9.3, 3])
para('따라서 신청·연장의 자격 기준은 “판정 시점의 최신 배치 스냅샷”입니다. 배치가 반드시 필요한 두 축은 '
     '① 기간 도래(규칙 발효·만료)와 ② 이 지표 재산정이며, 그 밖의 신청·가입·휴면복귀·상대형 종료는 배치 없이 '
     '즉시 또는 조회 시점에 반영됩니다.')
b('무효화 연동 :: 지표 재산정으로 자격이 바뀌면 해당 계좌의 협의 자격·해석 캐시를 무효화 → 다음 판정·조회에 반영')
b('실시간 필요 시 :: 최신성이 중요한 건은 신청 직전 온디맨드 재산정 훅을 둘 수 있으나, 기본은 비용·정합성 균형을 위해 배치 스냅샷을 참조')

doc.add_page_break()

# ============================================================ 10 캐시
h('10. 캐시와 무효화', level=1)
para('해석 결과는 (계좌, feeKey)로 RESOLVED_CACHE에 저장됩니다. read-through 방식이라 미스일 때만 계산하고 저장합니다.')
table(['동작', '설명'], [
    ['첫 조회(미스)', '후보 수집·게이트·최저가 계산 → 결과 저장'],
    ['재조회(적중)', '②③④ 및 최저가 계산을 건너뛰고 저장된 한 행을 바로 사용'],
    ['무효화', '규칙 승인/연장·지표 변경·협의 부여/해지 등으로 답이 달라질 수 있으면 영향 범위(scope 또는 계좌)만 무효화 → 다음 조회에서 lazy 재해석'],
], widths=[3, 12.5])
para('전량 재계산을 하지 않고 증분 무효화만 하므로, 규칙 변경이 드문 정상 운영에서는 무효화 빈도도 낮습니다. '
     '활성 계좌×빈출 조회키를 미리 데워 두는 캐시 예열을 쓰면 미스는 더욱 드물어집니다.')
doc.add_page_break()

# ============================================================ 11 등록/검증/승인
h('11. 등록 → 검증 → 승인 → 활성', level=1)
para('정책은 현업 기안 → 시뮬레이션 검증 → 심사 승인 → 활성의 흐름으로 반영됩니다. 활성 이후에는 계좌별 전개 없이 '
     '체결 시 해석됩니다.')
pic('d1_process')
num([
    '기안 :: 위저드에서 적용범위(상품군·거래소·세션·채널, 파생은 품목)·요율표·기간(2축)·대상을 지정',
    '지배관계 검증 :: 우대 요율표가 기준선을 공동 probe grid의 전 구간에서 하회하는지 자동 점검(불통과 구간이 있으면 그 구간에서 선택되지 않음)',
    '역마진 경고 :: 회사부담 합계가 자사 수취분을 초과하면 경고',
    '승인 :: 심사자 결재 → status=활성',
    '활성 후처리 :: 정책 우선순위 인덱스 갱신 + 영향 범위 캐시 무효화',
])
para('주식은 종목 단위로 우대를 걸지 않으므로 시뮬레이션·적용범위도 “거래소×세션×채널” 구간 단위로 다룹니다. '
     '종목별 지정은 파생상품에만 해당합니다.')

h('11.1 시스템 업무 흐름(현업·심사·플랫폼·원장)', level=2)
para('등록~활성(좌측)과 체결~부과(우측)를 주체별 스윔레인으로 보면 아래와 같습니다. 활성 이후에는 계좌별 전개 없이 '
     '체결 시점에 해석·부과가 일어납니다.')
pic('d9_swimlane', w=16.8)

out = '/Users/yujin-an/dev/fees/docs/수수료정책플랫폼_기술설계서_v0.8.docx'
doc.save(out); print('saved', out)
